from __future__ import annotations

import io
import json
import sys
import time
import zipfile
from pathlib import Path

from docx import Document

from webapp.app import app, _docx_preview_html, _write_job_artifacts_zip
from webapp.pdf_service import (
    Artifact,
    ConversionJobManager,
    ConversionResult,
    ConversionSubmission,
    ConversionOptions,
    LLM_MODEL_OPTIONS,
    NormalizedBlock,
    OPENROUTER_CHAT_COMPLETIONS_URL,
    PDFConversionService,
    _format_exam_blocks,
    _llm_api_key_env,
    _llm_model_for_provider,
    _llm_provider_for_model,
    _mineru_cli_from_python,
    _split_math_segments,
)


def test_content_list_v2_generates_editable_docx(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    data = [
        [
            {
                "type": "title",
                "content": {"title_content": [{"type": "text", "content": "1 Introduction"}], "level": 1},
            },
            {
                "type": "paragraph",
                "content": {"paragraph_content": [{"type": "text", "content": "A paragraph from MinerU."}]},
            },
            {
                "type": "table",
                "content": {
                    "table_caption": ["Table 1"],
                    "table_body": "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>",
                },
            },
        ]
    ]

    blocks, page_count = service._normalize_content_list_v2(data)
    output_path = tmp_path / "out.docx"
    service._write_docx(blocks, output_path, base_dirs=[tmp_path])

    doc = Document(output_path)
    assert page_count == 1
    assert [block.kind for block in blocks] == ["title", "paragraph", "table"]
    assert "1 Introduction" in [paragraph.text for paragraph in doc.paragraphs]
    assert doc.tables[0].cell(1, 1).text == "2"


def test_legacy_content_list_normalizes_core_blocks(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    data = [
        {"type": "text", "text": "Heading", "text_level": 2, "page_idx": 0},
        {"type": "list", "list_items": ["First", "Second"], "page_idx": 0},
        {"type": "equation", "text": "$$a=b$$", "page_idx": 1},
        {"type": "code", "code_body": "print('ok')", "page_idx": 1},
    ]

    blocks, page_count = service._normalize_content_list_legacy(data)

    assert page_count == 2
    assert blocks[0] == NormalizedBlock(kind="title", text="Heading", level=2, page_idx=0)
    assert blocks[1].kind == "list"
    assert blocks[1].items == ["First", "Second"]
    assert blocks[2].kind == "equation"
    assert blocks[3].kind == "code"


def test_nested_text_metadata_is_not_rendered_as_content(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    data = [
        {
            "type": "list",
            "list_items": [
                {
                    "type": "text",
                    "bbox": [1, 2, 3, 4],
                    "lines": [
                        {
                            "bbox": [1, 2, 3, 4],
                            "spans": [
                                {
                                    "type": "text",
                                    "bbox": [1, 2, 3, 4],
                                    "content": "a) Tại thời điểm bắt đầu phát hành video.",
                                }
                            ],
                        }
                    ],
                }
            ],
            "page_idx": 0,
        }
    ]

    blocks, _ = service._normalize_content_list_legacy(data)

    assert blocks[0].items == ["a) Tại thời điểm bắt đầu phát hành video."]


def test_v2_list_item_content_does_not_render_item_type(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    data = [
        [
            {
                "type": "list",
                "content": {
                    "list_type": "text_list",
                    "list_items": [
                        {
                            "item_type": "text",
                            "item_content": [
                                {
                                    "type": "text",
                                    "content": "a) Tại thời điểm bắt đầu phát hành video.",
                                }
                            ],
                        }
                    ],
                },
            }
        ]
    ]

    blocks, _ = service._normalize_content_list_v2(data)

    assert blocks[0].items == ["a) Tại thời điểm bắt đầu phát hành video."]


def test_explicit_lettered_list_items_are_not_bulleted(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    output_path = tmp_path / "lettered.docx"

    service._write_docx(
        [NormalizedBlock(kind="list", items=["a) First option", "b) Second option"])],
        output_path,
        base_dirs=[tmp_path],
    )

    doc = Document(output_path)
    assert [paragraph.text for paragraph in doc.paragraphs] == ["a) First option", "b) Second option"]
    assert all(paragraph.style.name != "List Bullet" for paragraph in doc.paragraphs)


def test_exam_options_are_spaced_sorted_and_wrapped_by_length() -> None:
    short = [
        NormalizedBlock(kind="paragraph", text="A. x=1."),
        NormalizedBlock(kind="paragraph", text="D.x=4."),
        NormalizedBlock(kind="paragraph", text="B.x=2."),
        NormalizedBlock(kind="paragraph", text="C.x=3."),
    ]
    medium = [
        NormalizedBlock(kind="paragraph", text="A. y = x + 1."),
        NormalizedBlock(kind="paragraph", text="B.y = x + 2026."),
        NormalizedBlock(kind="paragraph", text="C.y = 2x + 2026."),
        NormalizedBlock(kind="paragraph", text="D.y = 3x + 2026."),
    ]
    long = [
        NormalizedBlock(kind="paragraph", text="A.This is a deliberately long option body that should stay alone."),
        NormalizedBlock(kind="paragraph", text="B.This is a deliberately long option body that should stay alone."),
        NormalizedBlock(kind="paragraph", text="C.This is a deliberately long option body that should stay alone."),
        NormalizedBlock(kind="paragraph", text="D.This is a deliberately long option body that should stay alone."),
    ]

    rows = [block.text for block in _format_exam_blocks([*short, *medium, *long])]

    assert rows[0] == "A. x=1.\tB. x=2.\tC. x=3.\tD. x=4."
    assert rows[1] == "A. y = x + 1.\t\tB. y = x + 2026."
    assert rows[2] == "C. y = 2x + 2026.\t\tD. y = 3x + 2026."
    assert rows[3:] == [
        "A. This is a deliberately long option body that should stay alone.",
        "B. This is a deliberately long option body that should stay alone.",
        "C. This is a deliberately long option body that should stay alone.",
        "D. This is a deliberately long option body that should stay alone.",
    ]


def test_markdown_fallback_keeps_headings_lists_tables_and_code(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    blocks = service._normalize_markdown(
        """# Title

Intro text.

- One
- Two

| A | B |
| --- | --- |
| 1 | 2 |

```
print("ok")
```
"""
    )

    assert [block.kind for block in blocks] == ["title", "paragraph", "list", "table", "code"]
    assert blocks[2].items == ["One", "Two"]
    assert "<table>" in blocks[3].table_html


def test_markdown_details_blocks_are_not_rendered_to_docx(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    blocks = service._normalize_markdown(
        """Text before.

![](images/example.jpg)

<details>
<summary>natural_image</summary>

Cross-sectional generated image description.
</details>

Text after.
"""
    )

    assert not any("details" in block.text.lower() for block in blocks)
    assert not any("natural_image" in block.text for block in blocks)
    assert not any("Cross-sectional" in block.text for block in blocks)

    output_path = tmp_path / "details.docx"
    service._write_docx(blocks, output_path, base_dirs=[tmp_path])

    with zipfile.ZipFile(output_path) as docx:
        document_xml = docx.read("word/document.xml").decode("utf-8")

    assert "natural_image" not in document_xml
    assert "Cross-sectional" not in document_xml


def test_image_content_descriptions_are_not_rendered_to_docx(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    blocks, _ = service._normalize_content_list_v2(
        [
            [
                {
                    "type": "image",
                    "sub_type": "natural_image",
                    "content": {
                        "image_source": {"path": "images/example.jpg"},
                        "content": "Cross-sectional generated image description.",
                        "image_caption": [],
                        "image_footnote": [],
                    },
                }
            ]
        ]
    )

    assert len(blocks) == 1
    assert blocks[0].kind == "image"
    assert blocks[0].text == ""

    output_path = tmp_path / "image_description.docx"
    service._write_docx(blocks, output_path, base_dirs=[tmp_path])

    with zipfile.ZipFile(output_path) as docx:
        document_xml = docx.read("word/document.xml").decode("utf-8")

    assert "Cross-sectional" not in document_xml


def test_docx_math_spacing_handles_missing_source_spaces(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    blocks = [
        NormalizedBlock(kind="paragraph", text="Alpha\\(x=1\\)beta. Next \\(y=2\\).Gamma"),
    ]

    output_path = tmp_path / "spacing.docx"
    service._write_docx(blocks, output_path, base_dirs=[tmp_path])

    with zipfile.ZipFile(output_path) as docx:
        document_xml = docx.read("word/document.xml").decode("utf-8")

    assert "beta. Next " in document_xml
    assert ". Gamma" in document_xml
    assert "Next </w:t></w:r><w:r><w:t" not in document_xml


def test_table_and_markdown_formatting_in_docx(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    # Test case: Table with math and markdown
    table_html = "<table><tr><td>**Bold** $x=1$</td><td>*Italic* \\(y=2\\)</td></tr></table>"
    blocks = [
        NormalizedBlock(kind="paragraph", text="Testing **bold** and *italic* and `code`."),
        NormalizedBlock(kind="table", table_html=table_html),
    ]

    output_path = tmp_path / "format_test.docx"
    service._write_docx(blocks, output_path, base_dirs=[tmp_path])

    doc = Document(output_path)
    # Verify paragraph formatting
    p0 = doc.paragraphs[0]
    # "Testing " (plain), "bold" (bold), " and " (plain), "italic" (italic), " and " (plain), "code" (code), "." (plain)
    # Note: re.split might produce empty strings if separators are at the edges
    runs = [r for r in p0.runs if r.text]
    assert any(r.text == "bold" and r.bold for r in runs)
    assert any(r.text == "italic" and r.italic for r in runs)
    assert any(r.text == "code" and r.font.name == "Consolas" for r in runs)

    # Verify table math
    with zipfile.ZipFile(output_path) as docx:
        document_xml = docx.read("word/document.xml").decode("utf-8")

    assert "<m:oMath" in document_xml
    assert "x=1" in document_xml or ">x<" in document_xml
    assert "y=2" in document_xml or ">y<" in document_xml


def test_markdown_latex_is_written_as_word_math(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    blocks = service._normalize_markdown(
        """# Math

Inline \\(a+b\\) text.

\\[c=d\\]
"""
    )
    output_path = tmp_path / "math.docx"
    service._write_docx(blocks, output_path, base_dirs=[tmp_path])

    with zipfile.ZipFile(output_path) as docx:
        document_xml = docx.read("word/document.xml").decode("utf-8")

    assert "<m:oMath" in document_xml
    assert ">a<" in document_xml
    assert ">b<" in document_xml
    assert ">c<" in document_xml
    assert ">d<" in document_xml
    assert "\\(a+b\\)" not in document_xml
    assert "\\[c=d\\]" not in document_xml


def test_docx_preview_keeps_word_math_text(tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    blocks = [
        NormalizedBlock(kind="paragraph", text="Inline \\(a+b\\) text."),
        NormalizedBlock(kind="table", table_html="<table><tr><td>Cell \\(x=1\\)</td></tr></table>"),
    ]
    output_path = tmp_path / "preview_math.docx"
    service._write_docx(blocks, output_path, base_dirs=[tmp_path])

    html = _docx_preview_html(output_path)

    assert "Inline a+b text." in html
    assert "Cell x=1" in html


def test_plain_dollar_text_is_not_treated_as_math() -> None:
    assert _split_math_segments("The price is $5 and $10.") == [(False, "The price is $5 and $10.", False)]
    assert _split_math_segments("Use $x^2 + y^2$ here.") == [
        (False, "Use ", False),
        (True, "x^2 + y^2", False),
        (False, " here.", False),
    ]


def test_openrouter_llm_model_mapping_and_request(monkeypatch, tmp_path: Path) -> None:
    model = "openrouter/google/gemma-4-26b-a4b-it:free"
    assert model in {value for value, _ in LLM_MODEL_OPTIONS}
    assert _llm_provider_for_model(model) == "openrouter"
    assert _llm_api_key_env("openrouter") == "OPENROUTER_API_KEY"
    assert _llm_model_for_provider(model) == "google/gemma-4-26b-a4b-it:free"

    captured = {}

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def read(self):
            return b'{"choices":[{"message":{"content":"{}","reasoning_details":[{"type":"reasoning.text","text":"checked"}]}}]}'

    def fake_urlopen(request, timeout):
        captured["url"] = request.full_url
        captured["payload"] = json.loads(request.data.decode("utf-8"))
        captured["authorization"] = request.get_header("Authorization")
        captured["timeout"] = timeout
        return FakeResponse()

    monkeypatch.setattr("webapp.pdf_service.urllib.request.urlopen", fake_urlopen)

    service = PDFConversionService(tmp_path)
    response = service._call_llm_chat_completion(
        [{"role": "user", "content": "ping"}],
        model=model,
        api_key="test-key",
        reasoning=True,
    )

    assert response["content"] == "{}"
    assert response["reasoning_details"] == [{"type": "reasoning.text", "text": "checked"}]
    assert captured["url"] == OPENROUTER_CHAT_COMPLETIONS_URL
    assert captured["payload"]["model"] == "google/gemma-4-26b-a4b-it:free"
    assert captured["payload"]["reasoning"] == {"enabled": True}
    assert captured["authorization"] == "Bearer test-key"
    assert captured["timeout"] == 120


def test_openrouter_llm_failure_falls_back_to_nvidia(monkeypatch, tmp_path: Path) -> None:
    monkeypatch.setenv("OPENROUTER_API_KEY", "openrouter-key")
    monkeypatch.setenv("NVIDIA_API_KEY", "nvidia-key")
    service = PDFConversionService(tmp_path)
    review_dir = tmp_path / "review"
    calls = []

    def fake_call(messages, *, model, api_key, reasoning=False):
        calls.append((model, api_key, reasoning))
        if model.startswith("openrouter/"):
            raise RuntimeError("OpenRouter LLM HTTP 429: rate limited")
        return {"content": '{"findings":[],"patches":[]}', "reasoning_details": None}

    monkeypatch.setattr(service, "_call_llm_chat_completion", fake_call)

    blocks, warnings = service._run_llm_review_layer(
        [NormalizedBlock(kind="paragraph", text="abc", page_idx=0)],
        review_dir,
        ConversionOptions(
            llm_mode="correct",
            llm_model="openrouter/google/gemma-4-26b-a4b-it:free",
            llm_reasoning=True,
        ),
    )

    assert blocks[0].text == "abc"
    assert calls == [
        ("openrouter/google/gemma-4-26b-a4b-it:free", "openrouter-key", True),
        ("google/gemma-3-27b-it", "nvidia-key", True),
    ]
    assert any("fallback" in warning for warning in warnings)
    fallback_events = json.loads((review_dir / "fallback_events.json").read_text(encoding="utf-8"))
    assert fallback_events[0]["provider"] == "openrouter"
    assert fallback_events[0]["fallback_provider"] == "nvidia"
    summary = json.loads((review_dir / "review_request_summary.json").read_text(encoding="utf-8"))
    assert summary["fallback_events"] == 1
    assert summary["used_models"] == [{"chunk_index": 0, "model": "google/gemma-3-27b-it", "provider": "nvidia"}]


def test_run_mineru_passes_cli_options_and_config(monkeypatch, tmp_path: Path) -> None:
    service = PDFConversionService(tmp_path)
    captured = {}

    def fake_command() -> list[str]:
        return ["mineru"]

    def fake_run(command, **kwargs):
        captured["command"] = command
        captured["env"] = kwargs["env"]

        class Completed:
            returncode = 0
            stdout = ""
            stderr = ""

        return Completed()

    monkeypatch.setattr(service, "_mineru_command", fake_command)
    monkeypatch.setattr("webapp.pdf_service.subprocess.run", fake_run)
    output_dir = tmp_path / "out"
    output_dir.mkdir()

    options = ConversionOptions(
        backend="pipeline",
        parse_method="ocr",
        language="latin",
        formula_enable=False,
        table_enable=True,
        start_page=2,
        end_page=4,
        server_url="http://engine.test/v1",
        latex_delimiters_type="b",
    )
    service._run_mineru(tmp_path / "input.pdf", output_dir, "pipeline", options)

    command = captured["command"]
    assert command[:7] == ["mineru", "-p", str(tmp_path / "input.pdf"), "-o", str(output_dir), "-b", "pipeline"]
    assert ["-m", "ocr"] == command[7:9]
    assert "-l" in command and "latin" in command
    assert "-s" in command and "2" in command
    assert "-e" in command and "4" in command
    assert "-u" in command and "http://engine.test/v1" in command
    assert captured["env"]["MINERU_FORMULA_ENABLE"] == "false"
    assert Path(captured["env"]["MINERU_TOOLS_CONFIG_JSON"]).exists()


def test_job_manager_reaches_completed_status(tmp_path: Path) -> None:
    class FakeService:
        def convert(self, submission: ConversionSubmission) -> ConversionResult:
            docx_path = tmp_path / "jobs" / submission.job_id / "docx" / "out.docx"
            docx_path.parent.mkdir(parents=True, exist_ok=True)
            Document().save(docx_path)
            artifact = Artifact("Word DOCX", docx_path, "docx/out.docx", "docx")
            return ConversionResult(
                job_id=submission.job_id,
                original_filename=submission.original_filename,
                docx_path=docx_path,
                output_dir=tmp_path,
                artifacts=[artifact],
                backend_used="pipeline",
                elapsed_seconds=0.01,
                page_count=1,
                source_kind="content_list",
                source_path=None,
            )

    manager = ConversionJobManager(FakeService())  # type: ignore[arg-type]
    submission = ConversionSubmission("abc123", "input.pdf", tmp_path / "input.pdf", 100)
    manager.enqueue(submission)

    snapshot = None
    deadline = time.time() + 3
    while time.time() < deadline:
        snapshot = manager.get_snapshot("abc123")
        if snapshot and snapshot["status"] == "completed":
            break
        time.sleep(0.05)

    assert snapshot is not None
    assert snapshot["status"] == "completed"
    assert snapshot["result"]["backend_used"] == "pipeline"


def test_api_rejects_non_pdf_upload() -> None:
    client = app.test_client()
    response = client.post(
        "/api/convert",
        data={"pdf": (io.BytesIO(b"not a pdf"), "input.txt")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 400
    assert response.get_json()["ok"] is False


def test_pdf_to_word_uses_llm_model_select() -> None:
    client = app.test_client()
    response = client.get("/")

    assert response.status_code == 200
    html = response.get_data(as_text=True)
    assert '<select id="llmModelInput" name="llm_model">' in html
    assert "openrouter/google/gemma-4-26b-a4b-it:free" in html
    assert 'list="llmModelOptions"' not in html


def test_artifacts_zip_is_scoped_to_runtime_job_and_excludes_docx(tmp_path: Path) -> None:
    job_id = "abc123"
    job_dir = tmp_path / "webapp" / "runtime" / "jobs" / job_id
    (job_dir / "mineru").mkdir(parents=True)
    (job_dir / "docx").mkdir()
    (job_dir / "mineru" / "result.md").write_text("ok", encoding="utf-8")
    (job_dir / "docx" / "result.docx").write_bytes(b"docx")
    zip_path = job_dir / "artifacts_without_docx.zip"

    _write_job_artifacts_zip(job_dir, zip_path)

    with zipfile.ZipFile(zip_path) as archive:
        names = sorted(archive.namelist())

    assert names == [f"runtime/jobs/{job_id}/mineru/result.md"]


def test_readiness_blocks_python_314_env(monkeypatch, tmp_path: Path) -> None:
    if sys.version_info < (3, 14):
        return
    monkeypatch.setenv("MINERU_PYTHON_EXE", sys.executable)
    service = PDFConversionService(tmp_path)

    readiness = service.readiness()

    assert readiness.ready is False
    assert "Python 3.14" in readiness.message


def test_mineru_cli_fallback_checks_path(monkeypatch, tmp_path: Path) -> None:
    bin_dir = tmp_path / "bin"
    bin_dir.mkdir()
    executable = bin_dir / ("mineru.exe" if sys.platform == "win32" else "mineru")
    executable.write_text("", encoding="utf-8")
    monkeypatch.setenv("PATH", str(bin_dir))

    assert _mineru_cli_from_python(tmp_path / "python") == executable
