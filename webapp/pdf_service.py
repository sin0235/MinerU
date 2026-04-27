from __future__ import annotations

import json
import os
import re
import shlex
import shutil
import subprocess
import time
import uuid
import html.entities
import copy
import difflib
import inspect
import urllib.error
import urllib.request
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from pathlib import Path
from queue import Empty, Queue
from threading import Lock, Thread
from typing import Any, Callable

from bs4 import BeautifulSoup
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Emu, Inches, Pt
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename


class ConversionError(RuntimeError):
    """Raised when a PDF conversion cannot be completed."""


ALLOWED_BACKENDS = {
    "auto",
    "pipeline",
    "hybrid-auto-engine",
    "vlm-auto-engine",
    "hybrid-http-client",
    "vlm-http-client",
}
ALLOWED_PARSE_METHODS = {"auto", "txt", "ocr"}
ALLOWED_LANGUAGES = {
    "ch",
    "ch_server",
    "ch_lite",
    "en",
    "korean",
    "japan",
    "chinese_cht",
    "ta",
    "te",
    "ka",
    "th",
    "el",
    "latin",
    "arabic",
    "east_slavic",
    "cyrillic",
    "devanagari",
}
ALLOWED_LATEX_DELIMITER_TYPES = {"a", "b", "all"}
ALLOWED_LLM_MODES = {"off", "review", "correct"}
DEFAULT_NVIDIA_LLM_MODEL = "google/gemma-3-27b-it"
NVIDIA_CHAT_COMPLETIONS_URL = "https://integrate.api.nvidia.com/v1/chat/completions"
SKIPPED_CONTENT_TYPES = {
    "header",
    "footer",
    "page_header",
    "page_footer",
    "page_number",
    "page_aside_text",
    "aside_text",
    "page_footnote",
    "seal",
}
ProgressCallback = Callable[[dict[str, Any]], None]


@dataclass(slots=True)
class ConversionOptions:
    backend: str = "auto"
    parse_method: str = "auto"
    language: str = "ch"
    formula_enable: bool = True
    table_enable: bool = True
    start_page: int = 0
    end_page: int | None = None
    server_url: str = ""
    latex_delimiters_type: str = "b"
    exam_format: bool = False
    llm_mode: str = "off"
    llm_model: str = DEFAULT_NVIDIA_LLM_MODEL

    def to_payload(self) -> dict[str, Any]:
        return {
            "backend": self.backend,
            "parse_method": self.parse_method,
            "language": self.language,
            "formula_enable": self.formula_enable,
            "table_enable": self.table_enable,
            "start_page": self.start_page,
            "end_page": self.end_page,
            "server_url": self.server_url,
            "latex_delimiters_type": self.latex_delimiters_type,
            "exam_format": self.exam_format,
            "llm_mode": self.llm_mode,
            "llm_model": self.llm_model,
        }


@dataclass(slots=True)
class ReadinessInfo:
    ready: bool
    message: str
    command: list[str]
    backend: str
    python_version: str | None = None
    warnings: list[str] = field(default_factory=list)

    def to_payload(self) -> dict[str, Any]:
        return {
            "ready": self.ready,
            "message": self.message,
            "command": self.command,
            "backend": self.backend,
            "python_version": self.python_version,
            "warnings": list(self.warnings),
        }


@dataclass(slots=True)
class ConversionSubmission:
    job_id: str
    original_filename: str
    input_path: Path
    input_size_bytes: int
    options: ConversionOptions = field(default_factory=ConversionOptions)


@dataclass(slots=True)
class Artifact:
    label: str
    path: Path
    relative_path: str
    kind: str

    def to_payload(self) -> dict[str, Any]:
        return {
            "label": self.label,
            "filename": self.path.name,
            "relative_path": self.relative_path.replace("\\", "/"),
            "kind": self.kind,
            "size_bytes": self.path.stat().st_size if self.path.exists() else 0,
        }


@dataclass(slots=True)
class ConversionResult:
    job_id: str
    original_filename: str
    docx_path: Path
    output_dir: Path
    artifacts: list[Artifact]
    backend_used: str
    elapsed_seconds: float
    page_count: int
    source_kind: str
    source_path: Path | None
    warnings: list[str] = field(default_factory=list)

    def to_payload(self) -> dict[str, Any]:
        return {
            "job_id": self.job_id,
            "original_filename": self.original_filename,
            "download_name": self.docx_path.name,
            "backend_used": self.backend_used,
            "elapsed_seconds": round(self.elapsed_seconds, 2),
            "page_count": self.page_count,
            "source_kind": self.source_kind,
            "source_file": self.source_path.name if self.source_path else "",
            "warnings": list(self.warnings),
            "artifacts": [artifact.to_payload() for artifact in self.artifacts],
        }


@dataclass(slots=True)
class NormalizedBlock:
    kind: str
    text: str = ""
    level: int = 0
    items: list[str] = field(default_factory=list)
    table_html: str = ""
    image_path: str = ""
    caption: str = ""
    footnote: str = ""
    language: str = ""
    page_idx: int | None = None
    bbox: list[float] = field(default_factory=list)
    rich_content: list[dict[str, Any]] = field(default_factory=list)


class PDFConversionService:
    def __init__(self, root: Path) -> None:
        self.root = Path(root)
        self.runtime_dir = self.root / "webapp" / "runtime"
        self.jobs_dir = self.runtime_dir / "jobs"
        self.jobs_dir.mkdir(parents=True, exist_ok=True)

        self.max_upload_mb = _env_int("PDF_WORD_MAX_UPLOAD_MB", default=128, minimum=1, maximum=2048)
        self.max_upload_bytes = self.max_upload_mb * 1024 * 1024
        self.keep_artifacts = _env_flag("PDF_WORD_KEEP_ARTIFACTS", default=True)
        self.model_source = (os.getenv("MINERU_MODEL_SOURCE") or "huggingface").strip() or "huggingface"
        self.vl_model_name = (
            os.getenv("MINERU_VL_MODEL_NAME") or "opendatalab/MinerU2.5-Pro-2604-1.2B"
        ).strip()
        self.api_url = (os.getenv("MINERU_API_URL") or "").strip()
        self.timeout_seconds = _env_int("MINERU_TIMEOUT_SECONDS", default=3600, minimum=60, maximum=24 * 3600)

    def create_submission(self, upload: FileStorage) -> ConversionSubmission:
        return self.create_submission_with_options(upload, ConversionOptions())

    def create_submission_with_options(
        self,
        upload: FileStorage,
        options: ConversionOptions | None = None,
    ) -> ConversionSubmission:
        if not upload or not upload.filename:
            raise ValueError("Can upload file PDF.")

        original_filename = upload.filename
        extension = Path(original_filename).suffix.lower()
        if extension != ".pdf":
            raise ValueError("Chi ho tro file PDF.")

        job_id = uuid.uuid4().hex
        job_dir = self.job_dir(job_id)
        input_dir = job_dir / "input"
        input_dir.mkdir(parents=True, exist_ok=True)

        safe_name = secure_filename(Path(original_filename).stem) or "document"
        input_path = input_dir / f"{safe_name}.pdf"
        upload.save(input_path)
        size_bytes = input_path.stat().st_size
        if size_bytes <= 0:
            shutil.rmtree(job_dir, ignore_errors=True)
            raise ValueError("File PDF rong.")
        if size_bytes > self.max_upload_bytes:
            shutil.rmtree(job_dir, ignore_errors=True)
            raise ValueError(f"File PDF qua lon. Gioi han hien tai la {self.max_upload_mb} MB.")

        return ConversionSubmission(
            job_id=job_id,
            original_filename=original_filename,
            input_path=input_path,
            input_size_bytes=size_bytes,
            options=options or ConversionOptions(),
        )

    def job_dir(self, job_id: str) -> Path:
        safe_id = re.sub(r"[^a-f0-9]", "", job_id.lower())
        if not safe_id:
            raise ConversionError("Job id khong hop le.")
        return self.jobs_dir / safe_id

    def readiness(self) -> ReadinessInfo:
        command = self._mineru_command()
        backend = self.resolve_backend()
        warnings: list[str] = []
        python_version = self._configured_python_version()

        if python_version and _version_tuple(python_version) >= (3, 14):
            return ReadinessInfo(
                ready=False,
                message=(
                    "MinerU khong ho tro Python 3.14 tren Windows. Hay cai Python 3.12, tao env rieng, "
                    "roi dat MINERU_PYTHON_EXE hoac MINERU_COMMAND tro toi env do."
                ),
                command=command,
                backend=backend,
                python_version=python_version,
                warnings=warnings,
            )

        try:
            completed = subprocess.run(
                [*command, "--help"],
                cwd=self.root,
                text=True,
                encoding="utf-8",
                errors="replace",
                capture_output=True,
                timeout=20,
            )
        except FileNotFoundError:
            return ReadinessInfo(
                ready=False,
                message=(
                    "Khong tim thay lenh MinerU. Cai Python 3.12 + mineru[all], sau do dat "
                    "MINERU_PYTHON_EXE hoac MINERU_COMMAND."
                ),
                command=command,
                backend=backend,
                python_version=python_version,
                warnings=warnings,
            )
        except subprocess.TimeoutExpired:
            return ReadinessInfo(
                ready=False,
                message="Lenh MinerU --help bi timeout. Kiem tra env MinerU hoac thu chay lenh trong terminal.",
                command=command,
                backend=backend,
                python_version=python_version,
                warnings=warnings,
            )

        if completed.returncode != 0:
            detail = _compact_process_error(completed.stderr or completed.stdout)
            return ReadinessInfo(
                ready=False,
                message=f"MinerU chua san sang: {detail}",
                command=command,
                backend=backend,
                python_version=python_version,
                warnings=warnings,
            )

        if backend == "pipeline":
            warnings.append("Dang dung backend pipeline CPU; do chinh xac co the thap hon MinerU VLM.")
        if not self.api_url and backend in {"vlm-http-client", "hybrid-http-client"}:
            warnings.append("Backend HTTP client can MINERU_API_URL de ket noi MinerU API.")

        return ReadinessInfo(
            ready=True,
            message="MinerU CLI san sang.",
            command=command,
            backend=backend,
            python_version=python_version,
            warnings=warnings,
        )

    def resolve_backend(self, requested: str | None = None) -> str:
        requested = (requested or os.getenv("PDF_WORD_BACKEND") or "auto").strip() or "auto"
        if requested not in ALLOWED_BACKENDS:
            return "pipeline"
        if requested != "auto":
            return requested
        if self.api_url:
            return "vlm-http-client"
        return "hybrid-auto-engine" if self._mineru_cuda_available() else "pipeline"

    def convert(self, submission: ConversionSubmission, progress_callback: ProgressCallback | None = None) -> ConversionResult:
        readiness = self.readiness()
        if not readiness.ready:
            raise ConversionError(readiness.message)

        started = time.perf_counter()
        job_dir = self.job_dir(submission.job_id)
        mineru_output_dir = job_dir / "mineru"
        docx_dir = job_dir / "docx"
        mineru_output_dir.mkdir(parents=True, exist_ok=True)
        docx_dir.mkdir(parents=True, exist_ok=True)

        backend = self.resolve_backend(submission.options.backend)
        self._report_progress(
            progress_callback,
            progress=8,
            stage="prepare",
            message="Dang chuan bi thu muc job va kiem tra MinerU.",
        )
        self._run_mineru(submission.input_path, mineru_output_dir, backend, submission.options, progress_callback=progress_callback)

        self._report_progress(
            progress_callback,
            progress=70,
            stage="normalize",
            message="MinerU da xong. Dang doc cau truc tai lieu.",
        )
        blocks, source_path, source_kind, page_count, warnings = self._load_normalized_blocks(mineru_output_dir)
        if not blocks:
            raise ConversionError("MinerU khong tra ve noi dung doc duoc de tao DOCX.")

        if submission.options.llm_mode != "off":
            self._report_progress(
                progress_callback,
                progress=80,
                stage="llm_review",
                message="Dang chay lop LLM review tren noi dung da trich xuat.",
            )
            blocks, llm_warnings = self._run_llm_review_layer(blocks, job_dir / "llm_review", submission.options)
            warnings.extend(llm_warnings)

        self._report_progress(
            progress_callback,
            progress=88,
            stage="docx",
            message="Dang tao DOCX editable.",
        )
        docx_path = docx_dir / f"{secure_filename(Path(submission.original_filename).stem) or 'document'}.docx"
        self._write_docx(
            blocks,
            docx_path,
            base_dirs=[mineru_output_dir, source_path.parent if source_path else mineru_output_dir, submission.input_path.parent],
            options=submission.options,
        )

        self._report_progress(
            progress_callback,
            progress=96,
            stage="artifacts",
            message="Dang gom artifact va log ket qua.",
        )
        artifacts = self._collect_artifacts(job_dir, docx_path)
        if not self.keep_artifacts:
            self._remove_non_download_artifacts(job_dir, artifacts)
            artifacts = self._collect_artifacts(job_dir, docx_path)

        return ConversionResult(
            job_id=submission.job_id,
            original_filename=submission.original_filename,
            docx_path=docx_path,
            output_dir=mineru_output_dir,
            artifacts=artifacts,
            backend_used=backend,
            elapsed_seconds=time.perf_counter() - started,
            page_count=page_count,
            source_kind=source_kind,
            source_path=source_path,
            warnings=[*readiness.warnings, *warnings],
        )

    def resolve_download(self, job_id: str, relative_path: str) -> Path:
        job_dir = self.job_dir(job_id).resolve()
        candidate = (job_dir / relative_path).resolve()
        if not _is_relative_to(candidate, job_dir) or not candidate.exists() or not candidate.is_file():
            raise FileNotFoundError(relative_path)
        return candidate

    @staticmethod
    def _report_progress(progress_callback: ProgressCallback | None, **event: Any) -> None:
        if progress_callback is not None:
            progress_callback(event)

    def _run_mineru(
        self,
        pdf_path: Path,
        output_dir: Path,
        backend: str,
        options: ConversionOptions,
        progress_callback: ProgressCallback | None = None,
    ) -> None:
        command = [
            *self._mineru_command(),
            "-p",
            str(pdf_path),
            "-o",
            str(output_dir),
            "-b",
            backend,
            "-m",
            options.parse_method,
            "-l",
            options.language,
            "-f",
            _cli_bool(options.formula_enable),
            "-t",
            _cli_bool(options.table_enable),
        ]
        if options.start_page > 0:
            command.extend(["-s", str(options.start_page)])
        if options.end_page is not None:
            command.extend(["-e", str(options.end_page)])
        if options.server_url:
            command.extend(["-u", options.server_url])
        if self.api_url:
            command.extend(["--api-url", self.api_url])

        env = os.environ.copy()
        env.setdefault("MINERU_MODEL_SOURCE", self.model_source)
        env["MINERU_FORMULA_ENABLE"] = _env_bool(options.formula_enable)
        env["MINERU_TABLE_ENABLE"] = _env_bool(options.table_enable)
        env["MINERU_TOOLS_CONFIG_JSON"] = str(_write_mineru_config(output_dir, options))
        if self.vl_model_name:
            env.setdefault("MINERU_VL_MODEL_NAME", self.vl_model_name)

        if progress_callback is not None:
            self._run_mineru_streaming(command, output_dir, env, progress_callback)
            return

        completed = subprocess.run(
            command,
            cwd=self.root,
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            timeout=self.timeout_seconds,
            env=env,
        )
        (output_dir / "mineru_stdout.log").write_text(completed.stdout or "", encoding="utf-8")
        (output_dir / "mineru_stderr.log").write_text(completed.stderr or "", encoding="utf-8")
        if completed.returncode != 0:
            detail = _compact_process_error(completed.stderr or completed.stdout)
            raise ConversionError(f"MinerU xu ly that bai: {detail}")

    def _run_mineru_streaming(
        self,
        command: list[str],
        output_dir: Path,
        env: dict[str, str],
        progress_callback: ProgressCallback,
    ) -> None:
        self._report_progress(
            progress_callback,
            progress=14,
            stage="mineru",
            message="Dang chay MinerU. Terminal se cap nhat theo stdout/stderr.",
            terminal=f"$ {shlex.join(command)}",
        )
        process = subprocess.Popen(
            command,
            cwd=self.root,
            text=True,
            encoding="utf-8",
            errors="replace",
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            env=env,
        )
        output_queue: Queue[tuple[str, str]] = Queue()
        stdout_lines: list[str] = []
        stderr_lines: list[str] = []
        started = time.perf_counter()
        last_heartbeat = started

        def pump(stream: Any, stream_name: str) -> None:
            try:
                for raw_line in iter(stream.readline, ""):
                    output_queue.put((stream_name, raw_line.rstrip("\r\n")))
            finally:
                stream.close()

        threads = [
            Thread(target=pump, args=(process.stdout, "stdout"), daemon=True),
            Thread(target=pump, args=(process.stderr, "stderr"), daemon=True),
        ]
        for thread in threads:
            thread.start()

        while True:
            if time.perf_counter() - started > self.timeout_seconds:
                process.kill()
                for thread in threads:
                    thread.join(timeout=1)
                (output_dir / "mineru_stdout.log").write_text("\n".join(stdout_lines), encoding="utf-8")
                (output_dir / "mineru_stderr.log").write_text("\n".join(stderr_lines), encoding="utf-8")
                raise ConversionError(f"MinerU xu ly qua thoi gian cho phep ({self.timeout_seconds}s).")

            try:
                stream_name, line = output_queue.get(timeout=0.2)
            except Empty:
                if process.poll() is not None and output_queue.empty():
                    break
                now = time.perf_counter()
                if now - last_heartbeat >= 8:
                    last_heartbeat = now
                    progress_callback(
                        {
                            "stage": "mineru",
                            "message": "MinerU van dang chay, dang doi log moi tu tien trinh.",
                            "terminal": f"[runtime] MinerU van dang chay sau {round(now - started)}s.",
                        }
                    )
                continue

            last_heartbeat = time.perf_counter()
            if stream_name == "stdout":
                stdout_lines.append(line)
            else:
                stderr_lines.append(line)

            inferred_progress = _infer_mineru_progress(line)
            event: dict[str, Any] = {
                "stage": "mineru",
                "message": "MinerU dang phan tich PDF.",
                "terminal": f"[{stream_name}] {line}",
            }
            if inferred_progress is not None:
                event["progress"] = inferred_progress
            progress_callback(event)

        for thread in threads:
            thread.join(timeout=1)

        while not output_queue.empty():
            stream_name, line = output_queue.get_nowait()
            if stream_name == "stdout":
                stdout_lines.append(line)
            else:
                stderr_lines.append(line)
            progress_callback({"stage": "mineru", "terminal": f"[{stream_name}] {line}"})

        returncode = process.wait(timeout=1)
        (output_dir / "mineru_stdout.log").write_text("\n".join(stdout_lines), encoding="utf-8")
        (output_dir / "mineru_stderr.log").write_text("\n".join(stderr_lines), encoding="utf-8")
        if returncode != 0:
            detail = _compact_process_error("\n".join(stderr_lines) or "\n".join(stdout_lines))
            raise ConversionError(f"MinerU xu ly that bai: {detail}")

    def _run_llm_review_layer(
        self,
        blocks: list[NormalizedBlock],
        review_dir: Path,
        options: ConversionOptions,
    ) -> tuple[list[NormalizedBlock], list[str]]:
        warnings: list[str] = []
        review_dir.mkdir(parents=True, exist_ok=True)
        api_key = (os.getenv("NVIDIA_API_KEY") or "").strip()
        if not api_key:
            warnings.append("LLM review bi bo qua vi chua cau hinh NVIDIA_API_KEY.")
            (review_dir / "review_report.md").write_text(
                "# LLM Review\n\n- Status: skipped\n- Reason: NVIDIA_API_KEY is not configured\n",
                encoding="utf-8",
            )
            (review_dir / "review_request_summary.json").write_text(
                json.dumps({"mode": options.llm_mode, "skipped": True, "reason": "missing NVIDIA_API_KEY"}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            return blocks, warnings

        chunks = _chunk_blocks_for_llm(blocks)
        findings: list[dict[str, Any]] = []
        proposed_patches: list[dict[str, Any]] = []
        errors: list[dict[str, Any]] = []

        for chunk in chunks:
            try:
                response = self._call_nvidia_chat_completion(
                    _build_llm_messages(chunk, mode=options.llm_mode),
                    model=options.llm_model or DEFAULT_NVIDIA_LLM_MODEL,
                    api_key=api_key,
                )
                parsed = _parse_llm_json_response(response)
                findings.extend(parsed.get("findings") if isinstance(parsed.get("findings"), list) else [])
                if options.llm_mode == "correct":
                    proposed_patches.extend(parsed.get("patches") if isinstance(parsed.get("patches"), list) else [])
            except Exception as exc:
                errors.append({"chunk_index": chunk["chunk_index"], "error": str(exc)})

        updated_blocks = blocks
        applied: list[dict[str, Any]] = []
        rejected: list[dict[str, Any]] = []
        if options.llm_mode == "correct" and proposed_patches:
            updated_blocks, applied, rejected = _apply_safe_llm_patches(blocks, proposed_patches)
            if rejected:
                warnings.append(f"LLM da tu choi {len(rejected)} patch khong an toan.")
            if applied:
                warnings.append(f"LLM da ap dung {len(applied)} patch van ban an toan.")

        (review_dir / "review_findings.json").write_text(json.dumps(findings, ensure_ascii=False, indent=2), encoding="utf-8")
        (review_dir / "proposed_patches.json").write_text(json.dumps(proposed_patches, ensure_ascii=False, indent=2), encoding="utf-8")
        (review_dir / "applied_patches.json").write_text(json.dumps(applied, ensure_ascii=False, indent=2), encoding="utf-8")
        (review_dir / "rejected_patches.json").write_text(json.dumps(rejected, ensure_ascii=False, indent=2), encoding="utf-8")
        if errors:
            (review_dir / "llm_errors.json").write_text(json.dumps(errors, ensure_ascii=False, indent=2), encoding="utf-8")
            warnings.append(f"LLM review co {len(errors)} chunk loi; xem llm_errors.json.")
        (review_dir / "review_report.md").write_text(_llm_review_report(findings, applied, rejected, errors), encoding="utf-8")
        (review_dir / "review_request_summary.json").write_text(
            json.dumps(
                {
                    "mode": options.llm_mode,
                    "model": options.llm_model or DEFAULT_NVIDIA_LLM_MODEL,
                    "chunk_count": len(chunks),
                    "block_count": len(blocks),
                    "findings": len(findings),
                    "proposed_patches": len(proposed_patches),
                    "applied_patches": len(applied),
                    "rejected_patches": len(rejected),
                    "errors": len(errors),
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        return updated_blocks, warnings

    def _call_nvidia_chat_completion(self, messages: list[dict[str, str]], *, model: str, api_key: str) -> str:
        payload = json.dumps(
            {
                "model": model,
                "messages": messages,
                "max_tokens": 2048,
                "temperature": 0.0,
                "top_p": 0.7,
                "stream": False,
            }
        ).encode("utf-8")
        request = urllib.request.Request(
            NVIDIA_CHAT_COMPLETIONS_URL,
            data=payload,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Accept": "application/json",
                "Content-Type": "application/json",
            },
            method="POST",
        )
        try:
            with urllib.request.urlopen(request, timeout=120) as response:
                data = json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")[:1000]
            raise RuntimeError(f"NVIDIA LLM HTTP {exc.code}: {detail}") from exc
        content = (((data.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()
        if not content:
            raise RuntimeError("NVIDIA LLM tra ve noi dung rong.")
        return content

    def _mineru_command(self) -> list[str]:
        raw_command = (os.getenv("MINERU_COMMAND") or "").strip()
        if raw_command:
            return _split_command(raw_command)

        python_exe = (os.getenv("MINERU_PYTHON_EXE") or "").strip()
        if python_exe:
            cli_path = _mineru_cli_from_python(Path(python_exe))
            if cli_path is not None:
                return [str(cli_path)]
            return [python_exe, "-m", "mineru"]

        return ["mineru"]

    def _configured_python_version(self) -> str | None:
        python_exe = (os.getenv("MINERU_PYTHON_EXE") or "").strip()
        if not python_exe:
            return None
        try:
            completed = subprocess.run(
                [python_exe, "-c", "import sys; print(f'{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}')"],
                text=True,
                encoding="utf-8",
                errors="replace",
                capture_output=True,
                timeout=10,
            )
        except Exception:
            return None
        if completed.returncode != 0:
            return None
        return (completed.stdout or "").strip() or None

    def _mineru_cuda_available(self) -> bool:
        python_exe = (os.getenv("MINERU_PYTHON_EXE") or "").strip()
        if python_exe:
            try:
                completed = subprocess.run(
                    [
                        python_exe,
                        "-c",
                        "import torch; print('1' if torch.cuda.is_available() else '0')",
                    ],
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    capture_output=True,
                    timeout=10,
                )
                return completed.returncode == 0 and completed.stdout.strip() == "1"
            except Exception:
                return False

        try:
            import torch

            return bool(torch.cuda.is_available())
        except Exception:
            return False

    def _load_normalized_blocks(self, output_dir: Path) -> tuple[list[NormalizedBlock], Path | None, str, int, list[str]]:
        warnings: list[str] = []
        v2_path = _latest_file(output_dir, "*_content_list_v2.json")
        if v2_path:
            try:
                data = json.loads(v2_path.read_text(encoding="utf-8"))
                blocks, page_count = self._normalize_content_list_v2(data)
                if blocks:
                    return blocks, v2_path, "content_list_v2", page_count, warnings
                warnings.append("content_list_v2.json ton tai nhung khong co block doc duoc.")
            except Exception as exc:
                warnings.append(f"Khong doc duoc content_list_v2.json: {exc}")

        legacy_path = _latest_file(output_dir, "*_content_list.json")
        if legacy_path:
            try:
                data = json.loads(legacy_path.read_text(encoding="utf-8"))
                blocks, page_count = self._normalize_content_list_legacy(data)
                if blocks:
                    return blocks, legacy_path, "content_list", page_count, warnings
                warnings.append("content_list.json ton tai nhung khong co block doc duoc.")
            except Exception as exc:
                warnings.append(f"Khong doc duoc content_list.json: {exc}")

        markdown_path = _latest_file(output_dir, "*.md")
        if markdown_path:
            text = markdown_path.read_text(encoding="utf-8", errors="replace")
            blocks = self._normalize_markdown(text)
            if blocks:
                warnings.append("Da fallback sang Markdown vi khong co structured JSON phu hop.")
                return blocks, markdown_path, "markdown", 0, warnings

        return [], None, "none", 0, warnings

    def _normalize_content_list_v2(self, data: Any) -> tuple[list[NormalizedBlock], int]:
        blocks: list[NormalizedBlock] = []
        pages = _as_pages(data)
        for page_idx, page_items in enumerate(pages):
            for item in sorted(page_items, key=_reading_order_key):
                if not isinstance(item, dict):
                    continue
                kind = str(item.get("type") or "").strip()
                if not kind or kind in SKIPPED_CONTENT_TYPES:
                    continue
                content = item.get("content") if isinstance(item.get("content"), dict) else {}
                normalized = self._normalize_v2_item(kind, content, item, page_idx)
                if normalized is not None:
                    blocks.append(normalized)
        return blocks, len(pages)

    def _normalize_v2_item(
        self,
        kind: str,
        content: dict[str, Any],
        raw_item: dict[str, Any],
        page_idx: int,
    ) -> NormalizedBlock | None:
        if kind == "title":
            return NormalizedBlock(
                kind="title",
                text=_rich_text_to_plain(content.get("title_content")),
                level=_clamp_heading_level(content.get("level")),
                page_idx=page_idx,
            )
        if kind == "paragraph":
            rich_content = _rich_segments(content.get("paragraph_content"))
            text = _rich_segments_to_text(rich_content)
            return NormalizedBlock(
                kind="paragraph",
                text=text,
                page_idx=page_idx,
                bbox=_bbox(raw_item),
                rich_content=rich_content,
            ) if text else None
        if kind in {"list", "index"}:
            items = _to_string_list(content.get("list_items"))
            if not items:
                text = _rich_text_to_plain(content)
                items = [text] if text else []
            return NormalizedBlock(kind="list", items=items, page_idx=page_idx) if items else None
        if kind in {"equation_interline", "equation"}:
            text = _rich_text_to_plain(_first(content, "math_content", "equation_content", "text", "content"))
            image_path = _image_source_path(content) or str(raw_item.get("img_path") or "")
            return NormalizedBlock(kind="equation", text=text, image_path=image_path, page_idx=page_idx, bbox=_bbox(raw_item)) if text else None
        if kind in {"image", "chart"}:
            caption = _rich_text_to_plain(_first(content, f"{kind}_caption", "caption"))
            footnote = _rich_text_to_plain(_first(content, f"{kind}_footnote", "footnote"))
            image_path = _image_source_path(content) or str(_first(content, "image_path", "img_path", "path") or raw_item.get("img_path") or "")
            extracted = _rich_text_to_plain(_first(content, f"{kind}_content", "content"))
            return NormalizedBlock(
                kind=kind,
                text=extracted,
                image_path=image_path,
                caption=caption,
                footnote=footnote,
                page_idx=page_idx,
                bbox=_bbox(raw_item),
            )
        if kind == "table":
            caption = _rich_text_to_plain(_first(content, "table_caption", "caption"))
            footnote = _rich_text_to_plain(_first(content, "table_footnote", "footnote"))
            table_html = str(_first(content, "table_body", "table_html", "html") or "")
            image_path = _image_source_path(content) or str(_first(content, "image_path", "img_path", "path") or raw_item.get("img_path") or "")
            fallback_text = _rich_text_to_plain(_first(content, "table_content", "content"))
            return NormalizedBlock(
                kind="table",
                text=fallback_text,
                table_html=table_html,
                image_path=image_path,
                caption=caption,
                footnote=footnote,
                page_idx=page_idx,
                bbox=_bbox(raw_item),
            )
        if kind in {"code", "algorithm"}:
            code = _rich_text_to_plain(_first(content, "code_content", "algorithm_content", "code_body", "content"))
            caption = _rich_text_to_plain(_first(content, "code_caption", "algorithm_caption", "caption"))
            footnote = _rich_text_to_plain(_first(content, "code_footnote", "algorithm_footnote", "footnote"))
            language = str(_first(content, "code_language", "language") or "")
            return NormalizedBlock(
                kind="code",
                text=code,
                caption=caption,
                footnote=footnote,
                language=language,
                page_idx=page_idx,
            ) if code else None

        text = _rich_text_to_plain(content or raw_item.get("content"))
        return NormalizedBlock(kind="paragraph", text=text, page_idx=page_idx) if text else None

    def _normalize_content_list_legacy(self, data: Any) -> tuple[list[NormalizedBlock], int]:
        items = data if isinstance(data, list) else []
        blocks: list[NormalizedBlock] = []
        max_page_idx = -1
        for item in items:
            if not isinstance(item, dict):
                continue
            raw_kind = str(item.get("type") or "").strip()
            if not raw_kind or raw_kind in SKIPPED_CONTENT_TYPES:
                continue
            page_idx = item.get("page_idx") if isinstance(item.get("page_idx"), int) else None
            if page_idx is not None:
                max_page_idx = max(max_page_idx, page_idx)

            if raw_kind == "text":
                text = _rich_text_to_plain(item.get("text"))
                level = _clamp_heading_level(item.get("text_level"))
                blocks.append(
                    NormalizedBlock(
                        kind="title" if level else "paragraph",
                        text=text,
                        level=level,
                        page_idx=page_idx,
                    )
                )
            elif raw_kind == "list":
                items_text = _to_string_list(item.get("list_items"))
                if not items_text:
                    text = _rich_text_to_plain(item.get("text"))
                    items_text = [line.strip() for line in text.splitlines() if line.strip()]
                if items_text:
                    blocks.append(NormalizedBlock(kind="list", items=items_text, page_idx=page_idx))
            elif raw_kind in {"image", "chart"}:
                blocks.append(
                    NormalizedBlock(
                        kind=raw_kind,
                        text=_rich_text_to_plain(item.get("content")),
                        image_path=str(item.get("img_path") or item.get("image_path") or ""),
                        caption=_rich_text_to_plain(item.get(f"{raw_kind}_caption")),
                        footnote=_rich_text_to_plain(item.get(f"{raw_kind}_footnote")),
                        page_idx=page_idx,
                    )
                )
            elif raw_kind == "table":
                blocks.append(
                    NormalizedBlock(
                        kind="table",
                        table_html=str(item.get("table_body") or ""),
                        image_path=str(item.get("img_path") or item.get("image_path") or ""),
                        caption=_rich_text_to_plain(item.get("table_caption")),
                        footnote=_rich_text_to_plain(item.get("table_footnote")),
                        text=_rich_text_to_plain(item.get("content")),
                        page_idx=page_idx,
                    )
                )
            elif raw_kind == "equation":
                text = _rich_text_to_plain(item.get("text") or item.get("content"))
                if text:
                    blocks.append(NormalizedBlock(kind="equation", text=text, page_idx=page_idx))
            elif raw_kind == "code":
                code = _rich_text_to_plain(item.get("code_body") or item.get("content"))
                if code:
                    blocks.append(
                        NormalizedBlock(
                            kind="code",
                            text=code,
                            caption=_rich_text_to_plain(item.get("code_caption")),
                            footnote=_rich_text_to_plain(item.get("code_footnote")),
                            page_idx=page_idx,
                        )
                    )

        page_count = max_page_idx + 1 if max_page_idx >= 0 else 0
        return [block for block in blocks if _block_has_content(block)], page_count

    def _normalize_markdown(self, markdown: str) -> list[NormalizedBlock]:
        blocks: list[NormalizedBlock] = []
        paragraph_lines: list[str] = []
        list_items: list[str] = []
        table_lines: list[str] = []
        html_table_lines: list[str] = []
        code_lines: list[str] = []
        math_lines: list[str] = []
        math_right_delimiter = ""
        in_code = False
        in_math = False
        in_html_table = False

        def flush_paragraph() -> None:
            nonlocal paragraph_lines
            text = " ".join(line.strip() for line in paragraph_lines if line.strip()).strip()
            if text:
                blocks.append(NormalizedBlock(kind="paragraph", text=text))
            paragraph_lines = []

        def flush_list() -> None:
            nonlocal list_items
            if list_items:
                blocks.append(NormalizedBlock(kind="list", items=list_items))
            list_items = []

        def flush_table() -> None:
            nonlocal table_lines
            if table_lines:
                blocks.append(NormalizedBlock(kind="table", table_html=_markdown_table_to_html(table_lines)))
            table_lines = []

        def flush_html_table() -> None:
            nonlocal html_table_lines
            html_table = "\n".join(html_table_lines).strip()
            if html_table:
                blocks.append(NormalizedBlock(kind="table", table_html=html_table))
            html_table_lines = []

        def flush_code() -> None:
            nonlocal code_lines
            code = "\n".join(code_lines).strip("\n")
            if code:
                blocks.append(NormalizedBlock(kind="code", text=code))
            code_lines = []

        def flush_math() -> None:
            nonlocal math_lines
            text = "\n".join(math_lines).strip()
            if text:
                blocks.append(NormalizedBlock(kind="equation", text=_strip_math_delimiters(text)))
            math_lines = []

        for line in markdown.splitlines():
            stripped = line.strip()
            if in_html_table:
                html_table_lines.append(line)
                if "</table>" in stripped.lower():
                    flush_html_table()
                    in_html_table = False
                continue
            if in_math:
                math_lines.append(line)
                if stripped.endswith(math_right_delimiter):
                    flush_math()
                    in_math = False
                    math_right_delimiter = ""
                continue
            if stripped in {"$$", "\\["}:
                flush_paragraph()
                flush_list()
                flush_table()
                math_lines = [line]
                math_right_delimiter = "$$" if stripped == "$$" else "\\]"
                in_math = True
                continue
            if stripped.startswith("```"):
                if in_code:
                    flush_code()
                    in_code = False
                else:
                    flush_paragraph()
                    flush_list()
                    flush_table()
                    in_code = True
                continue
            if in_code:
                code_lines.append(line)
                continue
            if not stripped:
                flush_paragraph()
                flush_list()
                flush_table()
                continue
            heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading:
                flush_paragraph()
                flush_list()
                flush_table()
                blocks.append(
                    NormalizedBlock(kind="title", text=heading.group(2).strip(), level=min(len(heading.group(1)), 6))
                )
                continue
            if _is_display_math_line(stripped):
                flush_paragraph()
                flush_list()
                flush_table()
                blocks.append(NormalizedBlock(kind="equation", text=_strip_math_delimiters(stripped)))
                continue
            image_match = re.fullmatch(r'!\[[^\]]*\]\(\s*<?([^)>]+)>?(?:\s+"[^"]*")?\s*\)', stripped)
            if image_match:
                flush_paragraph()
                flush_list()
                flush_table()
                blocks.append(NormalizedBlock(kind="image", image_path=image_match.group(1).strip()))
                continue
            html_image_match = re.fullmatch(r"<img\b[^>]*\bsrc=[\"']([^\"']+)[\"'][^>]*>", stripped, flags=re.IGNORECASE)
            if html_image_match:
                flush_paragraph()
                flush_list()
                flush_table()
                blocks.append(NormalizedBlock(kind="image", image_path=html_image_match.group(1).strip()))
                continue
            if stripped.lower().startswith("<table"):
                flush_paragraph()
                flush_list()
                flush_table()
                html_table_lines = [line]
                if "</table>" in stripped.lower():
                    flush_html_table()
                else:
                    in_html_table = True
                continue
            list_match = re.match(r"^([-*+]|\d+[.)])\s+(.+)$", stripped)
            if list_match:
                flush_paragraph()
                flush_table()
                list_items.append(list_match.group(2).strip())
                continue
            if stripped.startswith("|") and stripped.endswith("|"):
                flush_paragraph()
                flush_list()
                table_lines.append(stripped)
                continue
            flush_list()
            flush_table()
            paragraph_lines.append(stripped)

        flush_code()
        flush_math()
        flush_html_table()
        flush_paragraph()
        flush_list()
        flush_table()
        return [block for block in blocks if _block_has_content(block)]

    def _write_docx(
        self,
        blocks: list[NormalizedBlock],
        output_path: Path,
        *,
        base_dirs: list[Path],
        options: ConversionOptions | None = None,
    ) -> None:
        options = options or ConversionOptions()
        document = Document()
        if options.exam_format:
            _apply_exam_section_format(document)
        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Times New Roman" if options.exam_format else "Arial"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), normal.font.name)
        normal.font.size = Pt(14 if options.exam_format else 11)

        for block in _format_exam_blocks(blocks) if options.exam_format else blocks:
            if block.kind == "title":
                document.add_heading(block.text, level=max(1, min(block.level or 1, 4)))
            elif block.kind == "paragraph":
                self._add_text_paragraph(document, block.text, rich_content=block.rich_content)
            elif block.kind == "list":
                for item in block.items:
                    if item.strip():
                        self._add_list_item(document, item.strip())
            elif block.kind == "table":
                self._add_table_block(document, block, base_dirs)
            elif block.kind in {"image", "chart"}:
                self._add_visual_block(document, block, base_dirs)
            elif block.kind == "equation":
                self._add_equation_block(document, block.text)
            elif block.kind == "code":
                if block.caption:
                    document.add_paragraph(block.caption).runs[0].italic = True
                paragraph = document.add_paragraph()
                run = paragraph.add_run(block.text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            if block.footnote:
                footnote = document.add_paragraph(block.footnote)
                if footnote.runs:
                    footnote.runs[0].italic = True

        if options.exam_format:
            _apply_exam_document_format(document)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    def _add_text_paragraph(self, document: Document, text: str, *, rich_content: list[dict[str, Any]] | None = None) -> None:
        paragraph = document.add_paragraph()
        if rich_content:
            _append_rich_content(paragraph, rich_content)
            return
        _append_text_with_math(paragraph, text)

    def _add_list_item(self, document: Document, text: str) -> None:
        if _has_explicit_list_marker(text):
            self._add_text_paragraph(document, text)
            return
        paragraph = document.add_paragraph(style="List Bullet")
        _append_text_with_math(paragraph, text)

    def _add_equation_block(self, document: Document, latex: str) -> None:
        paragraph = document.add_paragraph()
        _append_math(paragraph, _strip_math_delimiters(latex), display=True)

    def _add_table_block(self, document: Document, block: NormalizedBlock, base_dirs: list[Path]) -> None:
        if block.caption:
            document.add_paragraph(block.caption).runs[0].italic = True

        matrix = _html_table_to_matrix(block.table_html)
        if matrix:
            col_count = max(len(row) for row in matrix)
            table = document.add_table(rows=len(matrix), cols=col_count)
            table.style = "Table Grid"
            for row_idx, row in enumerate(matrix):
                for col_idx in range(col_count):
                    cell = table.cell(row_idx, col_idx)
                    text_content = row[col_idx] if col_idx < len(row) else ""
                    # Xoa paragraph mac dinh va dung _append_text_with_math de ho tro LaTeX
                    if cell.paragraphs:
                        p = cell.paragraphs[0]
                        p.clear()
                    else:
                        p = cell.add_paragraph()
                    _append_text_with_math(p, text_content)
            document.add_paragraph()
            return

        if block.text:
            self._add_text_paragraph(document, block.text)
            return

        self._add_visual_block(document, block, base_dirs)

    def _add_visual_block(self, document: Document, block: NormalizedBlock, base_dirs: list[Path]) -> None:
        if block.caption:
            document.add_paragraph(block.caption).runs[0].italic = True

        image_path = _resolve_artifact_path(block.image_path, base_dirs)
        if image_path and image_path.exists():
            try:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                width, height = _fit_image_size(image_path, document)
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=width, height=height)
            except Exception:
                document.add_paragraph(f"[Khong the chen anh: {image_path.name}]")
        elif block.image_path:
            paragraph = document.add_paragraph(f"[Không tìm thấy ảnh: {block.image_path}]")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if block.text:
            self._add_text_paragraph(document, block.text)

    def _collect_artifacts(self, job_dir: Path, docx_path: Path) -> list[Artifact]:
        candidates: list[Path] = [docx_path]
        patterns = [
            "*.md",
            "*_content_list_v2.json",
            "*_content_list.json",
            "*_layout.pdf",
            "*_span.pdf",
            "*_middle.json",
            "*_model.json",
            "mineru_stdout.log",
            "mineru_stderr.log",
            "llm_review/*.json",
            "llm_review/*.md",
        ]
        for pattern in patterns:
            candidates.extend(job_dir.rglob(pattern))

        artifacts: list[Artifact] = []
        seen: set[Path] = set()
        for path in candidates:
            if not path.exists() or not path.is_file():
                continue
            resolved = path.resolve()
            if resolved in seen:
                continue
            seen.add(resolved)
            artifacts.append(
                Artifact(
                    label=_artifact_label(path),
                    path=path,
                    relative_path=str(path.resolve().relative_to(job_dir.resolve())),
                    kind=_artifact_kind(path),
                )
            )
        return artifacts

    def _remove_non_download_artifacts(self, job_dir: Path, artifacts: list[Artifact]) -> None:
        keep = {artifact.path.resolve() for artifact in artifacts if artifact.kind in {"docx", "markdown", "json", "layout", "llm_review"}}
        for path in job_dir.rglob("*"):
            if path.is_file() and path.resolve() not in keep:
                path.unlink(missing_ok=True)


class ConversionJobManager:
    def __init__(self, service: PDFConversionService, *, max_workers: int = 1, retention_seconds: int = 24 * 60 * 60) -> None:
        self._service = service
        self._executor = ThreadPoolExecutor(max_workers=max_workers, thread_name_prefix="pdf-word-job")
        self._retention_seconds = retention_seconds
        self._jobs: dict[str, dict[str, Any]] = {}
        self._lock = Lock()

    def enqueue(self, submission: ConversionSubmission) -> str:
        self._purge_expired()
        now = time.time()
        snapshot = {
            "id": submission.job_id,
            "status": "queued",
            "created_at": now,
            "updated_at": now,
            "message": "Da nhan file PDF. Job dang cho worker xu ly.",
            "stage": "queued",
            "progress": 3,
            "terminal_lines": [],
            "original_filename": submission.original_filename,
            "input_size_bytes": submission.input_size_bytes,
            "result": None,
            "error": None,
        }
        with self._lock:
            self._jobs[submission.job_id] = snapshot
        self._executor.submit(self._run_job, submission)
        return submission.job_id

    def get_snapshot(self, job_id: str) -> dict[str, Any] | None:
        self._purge_expired()
        with self._lock:
            snapshot = self._jobs.get(job_id)
            return self._with_timing(snapshot) if snapshot else None

    def recent_results(self, *, limit: int = 8) -> list[dict[str, Any]]:
        self._purge_expired()
        with self._lock:
            completed = [
                dict(snapshot)
                for snapshot in self._jobs.values()
                if snapshot.get("status") == "completed" and isinstance(snapshot.get("result"), dict)
            ]
        completed.sort(key=lambda item: float(item.get("updated_at", 0)), reverse=True)
        return [self._with_timing(item) for item in completed[:limit]]

    @staticmethod
    def _with_timing(snapshot: dict[str, Any]) -> dict[str, Any]:
        item = dict(snapshot)
        now = time.time()
        created_at = float(item.get("created_at") or now)
        updated_at = float(item.get("updated_at") or created_at)
        result = item.get("result") if isinstance(item.get("result"), dict) else {}
        elapsed_seconds = float(result.get("elapsed_seconds") or max(0.0, now - created_at))
        item["elapsed_seconds"] = round(elapsed_seconds, 2)
        item["updated_age_seconds"] = round(max(0.0, now - updated_at), 2)
        progress = int(item.get("progress") or 0)
        if item.get("status") == "completed":
            item["eta_seconds"] = 0
        elif 5 <= progress < 100:
            eta = elapsed_seconds * (100 - progress) / progress
            item["eta_seconds"] = round(min(max(0.0, eta), 24 * 60 * 60), 2)
        else:
            item["eta_seconds"] = None
        return item

    def _run_job(self, submission: ConversionSubmission) -> None:
        self._update_job(
            submission.job_id,
            status="running",
            stage="starting",
            progress=6,
            message="MinerU dang phan tich PDF va tao file Word.",
        )
        progress_callback = self._progress_callback(submission.job_id)
        try:
            convert = self._service.convert
            try:
                supports_progress = "progress_callback" in inspect.signature(convert).parameters
            except (TypeError, ValueError):
                supports_progress = isinstance(self._service, PDFConversionService)
            if supports_progress:
                result = convert(submission, progress_callback=progress_callback)
            else:
                result = convert(submission)
        except ConversionError as exc:
            self._update_job(submission.job_id, status="failed", error=str(exc), message=str(exc))
            return
        except Exception as exc:
            self._update_job(
                submission.job_id,
                status="failed",
                error=f"Loi khong mong muon: {exc}",
                message=f"Loi khong mong muon: {exc}",
            )
            return

        self._update_job(
            submission.job_id,
            status="completed",
            stage="completed",
            progress=100,
            message="Da tao DOCX thanh cong.",
            result=result.to_payload(),
            error=None,
        )

    def _progress_callback(self, job_id: str) -> ProgressCallback:
        def callback(event: dict[str, Any]) -> None:
            self._apply_progress_event(job_id, event)

        return callback

    def _apply_progress_event(self, job_id: str, event: dict[str, Any]) -> None:
        with self._lock:
            snapshot = self._jobs.get(job_id)
            if snapshot is None:
                return
            current_progress = int(snapshot.get("progress") or 0)
            if "progress" in event:
                try:
                    snapshot["progress"] = max(current_progress, max(0, min(99, int(event["progress"]))))
                except (TypeError, ValueError):
                    pass
            if event.get("stage"):
                snapshot["stage"] = str(event["stage"])
            if event.get("message"):
                snapshot["message"] = str(event["message"])
            terminal_lines = snapshot.setdefault("terminal_lines", [])
            for raw_line in _event_terminal_lines(event):
                line = raw_line.strip()
                if line:
                    terminal_lines.append(line[:4000])
            snapshot["terminal_lines"] = terminal_lines[-400:]
            snapshot["updated_at"] = time.time()

    def _update_job(self, job_id: str, **updates: Any) -> None:
        with self._lock:
            snapshot = self._jobs.get(job_id)
            if snapshot is None:
                return
            snapshot.update(updates)
            snapshot["updated_at"] = time.time()

    def _purge_expired(self) -> None:
        now = time.time()
        with self._lock:
            expired = [
                job_id
                for job_id, snapshot in self._jobs.items()
                if now - float(snapshot.get("updated_at", snapshot.get("created_at", now))) > self._retention_seconds
            ]
            for job_id in expired:
                self._jobs.pop(job_id, None)


def _env_flag(name: str, *, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in {"1", "true", "yes", "on"}


def _env_int(name: str, *, default: int, minimum: int, maximum: int) -> int:
    raw = os.getenv(name)
    try:
        value = int(raw) if raw is not None else default
    except ValueError:
        value = default
    return max(minimum, min(maximum, value))


def _cli_bool(value: bool) -> str:
    return "true" if value else "false"


def _env_bool(value: bool) -> str:
    return "true" if value else "false"


def _write_mineru_config(output_dir: Path, options: ConversionOptions) -> Path:
    config: dict[str, Any] = {}
    configured_path = (os.getenv("MINERU_TOOLS_CONFIG_JSON") or "").strip()
    candidates = [Path(configured_path)] if configured_path else [Path.home() / "mineru.json"]
    for candidate in candidates:
        try:
            if candidate.exists():
                loaded = json.loads(candidate.read_text(encoding="utf-8"))
                if isinstance(loaded, dict):
                    config = loaded
                break
        except Exception:
            config = {}

    config["latex-delimiter-config"] = _latex_delimiter_config(options.latex_delimiters_type)
    path = output_dir / "mineru_config.json"
    path.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _latex_delimiter_config(delimiter_type: str) -> dict[str, dict[str, str]]:
    if delimiter_type == "a":
        return {
            "display": {"left": "$$", "right": "$$"},
            "inline": {"left": "$", "right": "$"},
        }
    if delimiter_type == "all":
        return {
            "display": {"left": "$$", "right": "$$"},
            "inline": {"left": "$", "right": "$"},
            "display_2": {"left": "\\[", "right": "\\]"},
            "inline_2": {"left": "\\(", "right": "\\)"},
        }
    return {
        "display": {"left": "\\[", "right": "\\]"},
        "inline": {"left": "\\(", "right": "\\)"},
    }


def _split_command(raw_command: str) -> list[str]:
    parts = shlex.split(raw_command, posix=os.name != "nt")
    return [part.strip('"') for part in parts if part.strip('"')]


def _mineru_cli_from_python(python_exe: Path) -> Path | None:
    candidates = [
        python_exe.parent / "mineru.exe",
        python_exe.parent / "mineru",
        python_exe.parent / "Scripts" / "mineru.exe",
        python_exe.parent / "Scripts" / "mineru",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    path_candidate = shutil.which("mineru")
    if path_candidate:
        return Path(path_candidate)
    return None


def _version_tuple(version: str) -> tuple[int, int, int]:
    parts = [int(match) for match in re.findall(r"\d+", version)[:3]]
    while len(parts) < 3:
        parts.append(0)
    return tuple(parts[:3])  # type: ignore[return-value]


def _compact_process_error(text: str, *, limit: int = 1200) -> str:
    compact = re.sub(r"\s+", " ", text or "").strip()
    if not compact:
        return "khong co stderr/stdout."
    return compact[:limit] + ("..." if len(compact) > limit else "")


def _infer_mineru_progress(line: str) -> int | None:
    text = line.strip()
    if not text:
        return None
    percent_match = re.search(r"(?<!\d)(\d{1,3})(?:\.\d+)?\s*%", text)
    if percent_match:
        raw_percent = max(0, min(100, int(percent_match.group(1))))
        return 15 + round(raw_percent * 0.5)
    page_match = re.search(r"(?:page|pages|trang)\D{0,12}(\d{1,5})\D{1,8}(\d{1,5})", text, flags=re.IGNORECASE)
    if page_match:
        current = int(page_match.group(1))
        total = max(1, int(page_match.group(2)))
        if current <= total:
            return 15 + round((current / total) * 50)
    lower = text.lower()
    phase_hints = [
        (18, ("load", "model", "init")),
        (28, ("parse", "analy", "ocr")),
        (42, ("layout", "detect")),
        (54, ("table", "formula", "span")),
        (64, ("dump", "save", "export")),
    ]
    for progress, words in phase_hints:
        if any(word in lower for word in words):
            return progress
    return None


def _event_terminal_lines(event: dict[str, Any]) -> list[str]:
    raw_lines = event.get("terminal_lines")
    lines: list[str] = []
    if isinstance(raw_lines, list):
        lines.extend(str(line) for line in raw_lines)
    raw_line = event.get("terminal")
    if raw_line is not None:
        lines.append(str(raw_line))
    return lines


def _as_pages(data: Any) -> list[list[dict[str, Any]]]:
    if not isinstance(data, list):
        return []
    if not data:
        return []
    if all(isinstance(item, list) for item in data):
        return [[entry for entry in page if isinstance(entry, dict)] for page in data]
    return [[entry for entry in data if isinstance(entry, dict)]]


def _first(mapping: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        value = mapping.get(key)
        if value not in (None, "", []):
            return value
    return None


def _rich_segments(value: Any) -> list[dict[str, str]]:
    if not isinstance(value, list):
        text = _rich_text_to_plain(value)
        return [{"type": "text", "content": text}] if text else []
    segments: list[dict[str, str]] = []
    for item in value:
        if isinstance(item, dict):
            kind = str(item.get("type") or "text")
            content = _rich_text_to_plain(item.get("content") if "content" in item else item)
            if content:
                segments.append({"type": kind, "content": content})
        else:
            content = _rich_text_to_plain(item)
            if content:
                segments.append({"type": "text", "content": content})
    return segments


def _rich_segments_to_text(segments: list[dict[str, str]]) -> str:
    pieces: list[tuple[bool, str]] = []
    for segment in segments:
        content = segment.get("content", "")
        if not content:
            continue
        if segment.get("type", "").startswith("equation"):
            pieces.append((True, f"\\({content}\\)"))
        else:
            pieces.append((False, content))
    return _join_text_math_pieces(pieces)


def _image_source_path(content: dict[str, Any]) -> str:
    image_source = content.get("image_source")
    if isinstance(image_source, dict):
        return str(image_source.get("path") or image_source.get("image_path") or image_source.get("img_path") or "")
    if isinstance(image_source, str):
        return image_source
    return ""


def _bbox(item: dict[str, Any]) -> list[float]:
    bbox = item.get("bbox")
    if isinstance(bbox, list):
        return [float(value) for value in bbox if isinstance(value, (int, float))]
    return []


def _reading_order_key(item: dict[str, Any]) -> tuple[float, float]:
    bbox = _bbox(item)
    if len(bbox) >= 2:
        return (bbox[1], bbox[0])
    return (0.0, 0.0)


def _rich_text_to_plain(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return re.sub(r"\s+", " ", value).strip()
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, list):
        pieces = [_rich_text_to_plain(item) for item in value]
        return re.sub(r"\s+", " ", " ".join(piece for piece in pieces if piece)).strip()
    if isinstance(value, dict):
        image_source = value.get("image_source")
        if isinstance(image_source, dict):
            path = _rich_text_to_plain(image_source.get("path"))
            if path:
                return path
        for key in (
            "content",
            "text",
            "paragraph_content",
            "title_content",
            "math_content",
            "table_content",
            "code_content",
            "code_body",
            "algorithm_content",
            "caption",
            "value",
            "spans",
            "lines",
            "blocks",
            "list_items",
            "item_content",
        ):
            if key in value:
                text = _rich_text_to_plain(value.get(key))
                if text:
                    return text
        metadata_keys = {
            "type",
            "sub_type",
            "bbox",
            "angle",
            "index",
            "score",
            "block_tags",
            "content_tags",
            "format",
            "page_idx",
            "page_size",
            "item_type",
            "list_type",
        }
        pieces = [_rich_text_to_plain(item) for key, item in value.items() if key not in metadata_keys]
        return re.sub(r"\s+", " ", " ".join(piece for piece in pieces if piece)).strip()
    return str(value).strip()


def _to_string_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [text for text in (_rich_text_to_plain(item) for item in value) if text]
    text = _rich_text_to_plain(value)
    return [text] if text else []


def _clamp_heading_level(value: Any) -> int:
    try:
        level = int(value)
    except (TypeError, ValueError):
        return 0
    return max(0, min(6, level))


def _block_has_content(block: NormalizedBlock) -> bool:
    return bool(block.text or block.items or block.table_html or block.image_path or block.caption)


def _chunk_blocks_for_llm(blocks: list[NormalizedBlock], *, max_chars: int = 6000) -> list[dict[str, Any]]:
    chunks: list[dict[str, Any]] = []
    current: list[dict[str, Any]] = []
    current_chars = 0
    current_page: int | None = None
    for index, block in enumerate(blocks):
        editable = _llm_editable_block(index, block)
        if editable is None:
            continue
        block_chars = len(json.dumps(editable, ensure_ascii=False))
        should_flush = current and (current_chars + block_chars > max_chars or editable.get("page_idx") != current_page)
        if should_flush:
            chunks.append({"chunk_index": len(chunks), "blocks": current})
            current = []
            current_chars = 0
        current.append(editable)
        current_chars += block_chars
        current_page = editable.get("page_idx")
    if current:
        chunks.append({"chunk_index": len(chunks), "blocks": current})
    return chunks


def _llm_editable_block(index: int, block: NormalizedBlock) -> dict[str, Any] | None:
    editable: dict[str, Any] = {"block_index": index, "page_idx": block.page_idx, "kind": block.kind, "fields": {}}
    if block.kind in {"paragraph", "title", "equation", "code"} and block.text:
        editable["fields"]["text"] = block.text
    if block.kind == "list" and block.items:
        editable["fields"]["items"] = block.items
    if block.caption:
        editable["fields"]["caption"] = block.caption
    if block.footnote:
        editable["fields"]["footnote"] = block.footnote
    return editable if editable["fields"] else None


def _build_llm_messages(chunk: dict[str, Any], *, mode: str) -> list[dict[str, str]]:
    action = "review and propose safe patches" if mode == "correct" else "review and report suspicious OCR issues"
    schema = (
        '{"findings":[{"block_index":0,"severity":"low|medium|high","issue_type":"ocr|math|spelling|format",'
        '"original":"...","suggestion":"...","reason":"..."}],'
        '"patches":[{"block_index":0,"field":"text|caption|footnote|items[0]","old_text":"...",'
        '"new_text":"...","confidence":0.0,"reason":"..."}]}'
    )
    return [
        {
            "role": "system",
            "content": (
                "You are a careful Vietnamese math exam OCR reviewer. Return strict JSON only. "
                "Do not translate or broadly rewrite. Do not change numbers, answer choices, formulas, names, paths, tables, or order unless the OCR error is obvious. "
                "For correct mode, patches must be small, high-confidence text fixes only."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Task: {action}. Output schema: {schema}\n\n"
                f"Chunk JSON:\n{json.dumps(chunk, ensure_ascii=False)}"
            ),
        },
    ]


def _parse_llm_json_response(raw: str) -> dict[str, Any]:
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    match = re.search(r"\{[\s\S]*\}", text)
    if match:
        text = match.group(0)
    parsed = json.loads(text)
    return parsed if isinstance(parsed, dict) else {}


def _apply_safe_llm_patches(
    blocks: list[NormalizedBlock], patches: list[dict[str, Any]]
) -> tuple[list[NormalizedBlock], list[dict[str, Any]], list[dict[str, Any]]]:
    updated = copy.deepcopy(blocks)
    applied: list[dict[str, Any]] = []
    rejected: list[dict[str, Any]] = []
    for patch in patches:
        ok, reason = _validate_llm_patch(updated, patch)
        if not ok:
            rejected.append({**patch, "rejected_reason": reason})
            continue
        block = updated[int(patch["block_index"])]
        field = str(patch["field"])
        old_text = str(patch["old_text"])
        new_text = str(patch["new_text"])
        if field == "text":
            block.text = block.text.replace(old_text, new_text, 1)
            if block.rich_content:
                for segment in block.rich_content:
                    if segment.get("type") == "text" and old_text in str(segment.get("content") or ""):
                        segment["content"] = str(segment.get("content") or "").replace(old_text, new_text, 1)
                        break
                else:
                    block.rich_content = []
        elif field == "caption":
            block.caption = block.caption.replace(old_text, new_text, 1)
        elif field == "footnote":
            block.footnote = block.footnote.replace(old_text, new_text, 1)
        else:
            match = re.fullmatch(r"items\[(\d+)\]", field)
            item_index = int(match.group(1))
            block.items[item_index] = block.items[item_index].replace(old_text, new_text, 1)
        applied.append(patch)
    return updated, applied, rejected


def _validate_llm_patch(blocks: list[NormalizedBlock], patch: dict[str, Any]) -> tuple[bool, str]:
    try:
        block_index = int(patch.get("block_index"))
    except (TypeError, ValueError):
        return False, "invalid block_index"
    if block_index < 0 or block_index >= len(blocks):
        return False, "block_index out of range"
    field = str(patch.get("field") or "")
    old_text = str(patch.get("old_text") or "")
    new_text = str(patch.get("new_text") or "")
    try:
        confidence = float(patch.get("confidence") or 0)
    except (TypeError, ValueError):
        return False, "invalid confidence"
    if confidence < 0.75:
        return False, "confidence too low"
    if not old_text or not new_text or old_text == new_text:
        return False, "empty or unchanged patch"
    block = blocks[block_index]
    target = ""
    if field in {"text", "caption", "footnote"}:
        target = getattr(block, field)
    else:
        match = re.fullmatch(r"items\[(\d+)\]", field)
        if not match:
            return False, "unsupported field"
        item_index = int(match.group(1))
        if item_index < 0 or item_index >= len(block.items):
            return False, "item index out of range"
        target = block.items[item_index]
    if old_text not in target:
        return False, "old_text mismatch"
    if _numbers_changed(old_text, new_text):
        return False, "numbers changed"
    if _latex_commands_changed(old_text, new_text):
        return False, "latex commands changed"
    ratio = difflib.SequenceMatcher(None, old_text, new_text).ratio()
    if ratio < 0.55 or abs(len(new_text) - len(old_text)) > max(20, len(old_text) * 0.4):
        return False, "change too large"
    return True, ""


def _numbers_changed(old_text: str, new_text: str) -> bool:
    return re.findall(r"\d+(?:[,.]\d+)?", old_text) != re.findall(r"\d+(?:[,.]\d+)?", new_text)


def _latex_commands_changed(old_text: str, new_text: str) -> bool:
    return re.findall(r"\\[A-Za-z]+", old_text) != re.findall(r"\\[A-Za-z]+", new_text)


def _llm_review_report(
    findings: list[dict[str, Any]], applied: list[dict[str, Any]], rejected: list[dict[str, Any]], errors: list[dict[str, Any]]
) -> str:
    lines = ["# LLM Review", "", f"- Findings: {len(findings)}", f"- Applied patches: {len(applied)}", f"- Rejected patches: {len(rejected)}", f"- Errors: {len(errors)}", ""]
    if findings:
        lines.append("## Findings")
        for item in findings:
            lines.append(f"- Block {item.get('block_index')}: {item.get('issue_type')} / {item.get('severity')} — {item.get('reason')}")
            if item.get("original") or item.get("suggestion"):
                lines.append(f"  - `{item.get('original', '')}` → `{item.get('suggestion', '')}`")
    if applied:
        lines.extend(["", "## Applied patches"])
        for patch in applied:
            lines.append(f"- Block {patch.get('block_index')} `{patch.get('field')}`: `{patch.get('old_text')}` → `{patch.get('new_text')}`")
    if rejected:
        lines.extend(["", "## Rejected patches"])
        for patch in rejected:
            lines.append(f"- Block {patch.get('block_index')} `{patch.get('field')}`: {patch.get('rejected_reason')}")
    if errors:
        lines.extend(["", "## Errors"])
        for error in errors:
            lines.append(f"- Chunk {error.get('chunk_index')}: {error.get('error')}")
    return "\n".join(lines).strip() + "\n"


def _format_exam_blocks(blocks: list[NormalizedBlock]) -> list[NormalizedBlock]:
    formatted: list[NormalizedBlock] = []
    option_buffer: list[NormalizedBlock] = []

    def flush_options() -> None:
        nonlocal option_buffer
        if option_buffer:
            formatted.extend(_layout_exam_options(option_buffer))
            option_buffer = []

    for block in blocks:
        if block.kind == "paragraph" and _is_exam_option(block.text):
            option_buffer.append(block)
            if len(option_buffer) == 4:
                flush_options()
            continue
        if block.kind == "list" and block.items and all(_is_exam_option(item) for item in block.items):
            flush_options()
            formatted.extend(_layout_exam_options([NormalizedBlock(kind="paragraph", text=item) for item in block.items]))
            continue
        flush_options()
        formatted.append(block)
    flush_options()
    return formatted


def _layout_exam_options(options: list[NormalizedBlock]) -> list[NormalizedBlock]:
    texts = [_normalize_exam_option_text(block.text) for block in _sort_exam_options(options)]
    max_len = max((_plain_length(_option_body(text)) for text in texts), default=0)
    if max_len > 38:
        per_line = 1
    elif max_len > 10:
        per_line = 2
    else:
        per_line = 4

    rows: list[NormalizedBlock] = []
    for index in range(0, len(texts), per_line):
        separator = "\t\t" if per_line == 2 else "\t"
        rows.append(NormalizedBlock(kind="paragraph", text=separator.join(texts[index : index + per_line])))
    return rows


def _is_exam_option(text: str) -> bool:
    return bool(_exam_option_match(text))


def _option_body(text: str) -> str:
    match = _exam_option_match(text)
    return text.strip()[match.end() :].strip() if match else text.strip()


def _has_explicit_list_marker(text: str) -> bool:
    return bool(re.match(r"^(?:[A-Za-z]|\d+)[.)]\s*", text.strip()))


def _exam_option_match(text: str) -> re.Match[str] | None:
    return re.match(r"^([A-Da-d])([.)])\s*", text.strip())


def _normalize_exam_option_text(text: str) -> str:
    stripped = text.strip()
    match = _exam_option_match(stripped)
    if not match:
        return stripped
    marker = match.group(1).upper() if match.group(1).isupper() else match.group(1).lower()
    punctuation = "." if marker.isupper() else match.group(2)
    body = stripped[match.end() :].strip()
    return f"{marker}{punctuation} {body}" if body else f"{marker}{punctuation}"


def _sort_exam_options(options: list[NormalizedBlock]) -> list[NormalizedBlock]:
    markers = [_exam_option_match(block.text) for block in options]
    if len(options) == 4 and all(markers):
        letters = [match.group(1).upper() for match in markers if match]
        if sorted(letters) == ["A", "B", "C", "D"]:
            return sorted(options, key=_exam_option_sort_key)
    return options


def _exam_option_sort_key(block: NormalizedBlock) -> str:
    match = _exam_option_match(block.text)
    return match.group(1).upper() if match else ""


def _plain_length(text: str) -> int:
    text = _strip_math_delimiters(text)
    text = re.sub(r"\\(?:frac|dfrac|tfrac)\s*\{([^{}]*)\}\s*\{([^{}]*)\}", r"\1/\2", text)
    text = re.sub(r"\\(?:vec|overrightarrow)\s*\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[A-Za-z]+", "", text)
    text = re.sub(r"[{}\\]", "", text)
    return len(text.strip())


def _apply_exam_document_format(document: Document) -> None:
    _apply_exam_section_format(document)
    for paragraph in document.paragraphs:
        _apply_exam_paragraph_format(paragraph)
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _apply_exam_paragraph_format(paragraph, in_table=True)


def _apply_exam_section_format(document: Document) -> None:
    for section in document.sections:
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def _apply_exam_paragraph_format(paragraph: Any, *, in_table: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if _paragraph_has_picture(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if in_table else WD_ALIGN_PARAGRAPH.JUSTIFY
    text = paragraph.text
    if "\t" in text:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        stops = fmt.tab_stops
        if "\t\t" in text:
            stops.add_tab_stop(Cm(8.2))
        else:
            stops.add_tab_stop(Cm(3.8))
            stops.add_tab_stop(Cm(7.6))
            stops.add_tab_stop(Cm(11.4))
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)
    _style_exam_labels(paragraph)


def _paragraph_has_picture(paragraph: Any) -> bool:
    return bool(paragraph._p.xpath(".//pic:pic"))


def _style_exam_labels(paragraph: Any) -> None:
    text = paragraph.text
    question_match = re.match(r"^(Câu\s+\d+\s*[.:)])", text, flags=re.IGNORECASE)
    if question_match:
        _style_text_prefix(paragraph, len(question_match.group(1)), bold=True, italic=True)
    for match in re.finditer(r"(^|\t)([A-Da-d][.)])\s*", text):
        _style_text_range(paragraph, match.start(2), match.end(2), bold=True)


def _style_text_prefix(paragraph: Any, length: int, *, bold: bool = False, italic: bool = False) -> None:
    _style_text_range(paragraph, 0, length, bold=bold, italic=italic)


def _style_text_range(paragraph: Any, start: int, end: int, *, bold: bool = False, italic: bool = False) -> None:
    cursor = 0
    for run in list(paragraph.runs):
        run_text = run.text or ""
        run_len = len(run_text)
        overlap_start = max(start, cursor)
        overlap_end = min(end, cursor + run_len)
        if overlap_start < overlap_end:
            local_start = overlap_start - cursor
            local_end = overlap_end - cursor
            if local_start == 0 and local_end == run_len:
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
            else:
                _split_and_style_run(run, local_start, local_end, bold=bold, italic=italic)
        cursor += run_len


def _split_and_style_run(run: Any, start: int, end: int, *, bold: bool = False, italic: bool = False) -> None:
    text = run.text or ""
    before, middle, after = text[:start], text[start:end], text[end:]
    run.text = before
    styled = run._parent.add_run(middle)
    _copy_run_style(run, styled)
    if bold:
        styled.bold = True
    if italic:
        styled.italic = True
    if after:
        trailing = run._parent.add_run(after)
        _copy_run_style(run, trailing)
    paragraph_element = run._parent._p
    base = run._r
    if after:
        paragraph_element.remove(trailing._r)
        base.addnext(trailing._r)
    paragraph_element.remove(styled._r)
    base.addnext(styled._r)


def _copy_run_style(source: Any, target: Any) -> None:
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.font.name = source.font.name
    target.font.size = source.font.size
    target._element.rPr.rFonts.set(qn("w:eastAsia"), source.font.name or "Times New Roman")


def _fit_image_size(image_path: Path, document: Document) -> tuple[Emu, Emu | None]:
    section = document.sections[-1]
    available_width = section.page_width - section.left_margin - section.right_margin
    max_width = min(available_width, Inches(6.8))
    max_height = Inches(4.8)

    try:
        from PIL import Image

        with Image.open(image_path) as image:
            width_px, height_px = image.size
            dpi_x, dpi_y = image.info.get("dpi", (96, 96))
    except Exception:
        return Emu(max_width), None

    if width_px <= 0 or height_px <= 0:
        return Emu(max_width), None

    dpi_x = dpi_x if isinstance(dpi_x, (int, float)) and dpi_x > 0 else 96
    dpi_y = dpi_y if isinstance(dpi_y, (int, float)) and dpi_y > 0 else 96
    native_width = Inches(width_px / dpi_x)
    native_height = Inches(height_px / dpi_y)
    scale = min(max_width / native_width, max_height / native_height, 1.0)
    width = max(1, int(native_width * scale))
    height = max(1, int(native_height * scale))
    return Emu(width), Emu(height)


def _append_formatted_text(paragraph: Any, text: str) -> None:
    # Parser don gian cho inline markdown: **bold**, *italic*, `code`
    # Dung regex de split text thanh cac segment co format
    pattern = r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|__.*?__|__.*?__|_.*?_|\`.*?\`)"
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if (part.startswith("***") and part.endswith("***")) or (part.startswith("___") and part.endswith("___")):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif (part.startswith("**") and part.endswith("**")) or (part.startswith("__") and part.endswith("__")):
            run.text = part[2:-2]
            run.bold = True
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run.text = part.replace("\\{", "{").replace("\\}", "}")


def _append_rich_content(paragraph: Any, segments: list[dict[str, Any]]) -> None:
    pending: list[tuple[bool, str, bool]] = []
    for segment in segments:
        kind = str(segment.get("type") or "text")
        content = str(segment.get("content") or "")
        if not content:
            continue
        if kind.startswith("equation"):
            pending.append((True, content, False))
        else:
            pending.extend(_expand_text_math_segments(content))
    _append_spaced_math_segments(paragraph, pending)


def _append_text_with_math(paragraph: Any, text: str) -> None:
    _append_spaced_math_segments(paragraph, _expand_text_math_segments(text))


def _expand_text_math_segments(text: str) -> list[tuple[bool, str, bool]]:
    segments: list[tuple[bool, str, bool]] = []
    for is_math, value, display in _split_math_segments(text):
        if not value:
            continue
        if is_math:
            segments.append((True, value, display))
            continue
        for implicit_is_math, implicit_value in _split_implicit_latex_segments(value):
            if implicit_value:
                segments.append((implicit_is_math, implicit_value, False))
    return segments


def _append_spaced_math_segments(paragraph: Any, segments: list[tuple[bool, str, bool]]) -> None:
    previous_value = ""
    previous_is_math = False
    for index, (is_math, value, display) in enumerate(segments):
        if not value:
            continue
        next_is_math = False
        next_value = ""
        for candidate_is_math, candidate_value, _ in segments[index + 1 :]:
            if candidate_value:
                next_is_math = candidate_is_math
                next_value = candidate_value
                break
        if is_math:
            if _needs_space_between(previous_value, value, previous_is_math, True):
                paragraph.add_run(" ")
            _append_math(paragraph, value, display=display)
            if _needs_space_between(value, next_value, True, next_is_math):
                paragraph.add_run(" ")
        else:
            _append_formatted_text(paragraph, value)
        previous_value = value
        previous_is_math = is_math


def _needs_space_between(left: str, right: str, left_is_math: bool, right_is_math: bool) -> bool:
    if not left or not right or left_is_math == right_is_math:
        return False
    left_edge = _visible_edge(left, from_right=True)
    right_edge = _visible_edge(right, from_right=False)
    if not left_edge or not right_edge:
        return False
    if left_edge.isspace() or right_edge.isspace():
        return False
    if left_edge in {'(', '[', '{', '/', '\\', '"', "'", '“', '‘'} or right_edge in {')', ']', '}', '/', ',', '.', ';', ':', '!', '?', '%', '”', '’'}:
        return False
    return _is_wordish_boundary(left_edge) or _is_wordish_boundary(right_edge)


def _visible_edge(text: str, *, from_right: bool) -> str:
    stripped = text.rstrip() if from_right else text.lstrip()
    return stripped[-1:] if from_right else stripped[:1]


def _is_wordish_boundary(char: str) -> bool:
    return char.isalnum() or char in "}_^'′" or ord(char) > 127


def _join_text_math_pieces(pieces: list[tuple[bool, str]]) -> str:
    output: list[str] = []
    previous_is_math = False
    previous_value = ""
    for is_math, value in pieces:
        if not value:
            continue
        if output and _needs_space_between(previous_value, value, previous_is_math, is_math):
            output.append(" ")
        output.append(value)
        previous_value = value
        previous_is_math = is_math
    return re.sub(r"\s+", " ", "".join(output)).strip()


def _append_math(paragraph: Any, latex: str, *, display: bool) -> None:
    latex = _normalize_latex(_strip_math_delimiters(latex))
    if not latex:
        return
    omml_xml = _latex_to_omml_xml(latex)
    if omml_xml:
        paragraph._p.append(parse_xml(_ensure_omml_namespaces(omml_xml)))
        return
    if _append_vector_latex_fallback(paragraph, latex):
        return
    paragraph.add_run(_latex_as_text(latex, display=display))


def _latex_to_omml_xml(latex: str) -> str | None:
    try:
        import latex2mathml.converter
        import mathml2omml
    except Exception:
        return None
    try:
        mathml = latex2mathml.converter.convert(latex)
        convert = getattr(mathml2omml, "convert", None)
        if convert is None:
            return None
        omml_xml = str(convert(mathml, html.entities.name2codepoint))
        parse_xml(_ensure_omml_namespaces(omml_xml))
        return omml_xml
    except Exception:
        return None


def _append_vector_latex_fallback(paragraph: Any, latex: str) -> bool:
    pattern = re.compile(r"\\(?:vec|overrightarrow)\s*\{(?:[^{}]|\{[^{}]*\})+\}")
    if not pattern.search(latex):
        return False
    cursor = 0
    appended = False
    for match in pattern.finditer(latex):
        if match.start() > cursor:
            _append_text_with_math(paragraph, _latex_as_text(latex[cursor:match.start()], display=False))
        vector_omml_xml = _vector_latex_to_omml_xml(match.group(0))
        if vector_omml_xml:
            paragraph._p.append(parse_xml(_ensure_omml_namespaces(vector_omml_xml)))
            appended = True
        else:
            paragraph.add_run(_latex_as_text(match.group(0), display=False))
        cursor = match.end()
    if cursor < len(latex):
        _append_text_with_math(paragraph, _latex_as_text(latex[cursor:], display=False))
    return appended


def _vector_latex_to_omml_xml(latex: str) -> str | None:
    match = re.fullmatch(r"\\(?:vec|overrightarrow)\s*\{(.+)\}", latex.strip())
    if not match:
        return None
    value = _latex_to_plain_math_text(match.group(1).strip())
    if not value:
        return None
    escaped_value = _escape_xml_text(value)
    return (
        '<m:oMath>'
        '<m:acc>'
        '<m:accPr><m:chr m:val="⃗"/></m:accPr>'
        f'<m:e><m:r><m:t>{escaped_value}</m:t></m:r></m:e>'
        '</m:acc>'
        '</m:oMath>'
    )


def _latex_to_plain_math_text(text: str) -> str:
    text = re.sub(r"\^\s*\{\s*\\prime\s*\}", "′", text)
    text = re.sub(r"\\prime\b", "′", text)
    text = re.sub(r"\^\s*\{([^{}]+)\}", r"^\1", text)
    text = re.sub(r"_\s*\{([^{}]+)\}", r"_\1", text)
    return text


def _escape_xml_text(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _ensure_omml_namespaces(omml_xml: str) -> str:
    if "xmlns:m=" in omml_xml:
        return omml_xml
    for root in ("m:oMathPara", "m:oMath"):
        prefix = f"<{root}"
        if omml_xml.startswith(prefix):
            return omml_xml.replace(prefix, f"{prefix} {nsdecls('m')}", 1)
    return omml_xml


def _normalize_latex(latex: str) -> str:
    text = latex.strip()
    text = re.sub(r"\\tag\s*\{[^{}]*\}", "", text)
    text = re.sub(r"\\label\s*\{[^{}]*\}", "", text)
    text = re.sub(r"\\(displaystyle|textstyle|scriptstyle|scriptscriptstyle)\b", "", text)
    text = re.sub(r"\\left\s*([.])", r"\\left\\1", text)
    text = re.sub(r"\\right\s*([.])", r"\\right\\1", text)
    text = text.replace("\\,", " ").replace("\\;", " ").replace("\\:", " ").replace("\\!", "")
    text = re.sub(r"\\operatorname\s*\{([^{}]+)\}", r"\\mathrm{\1}", text)
    text = re.sub(r"\\text\s*\{([^{}]*)\}", r"\\mathrm{\1}", text)

    environments = ("aligned", "align", "align*", "gathered", "gather", "gather*", "split")
    for environment in environments:
        text = text.replace(f"\\begin{{{environment}}}", r"\\begin{array}{l}")
        text = text.replace(f"\\end{{{environment}}}", r"\\end{array}")

    if "\\begin{cases}" in text and "\\begin{array}" not in text:
        text = text.replace("\\begin{cases}", "\\begin{cases} \\begin{array}{l}")
        text = text.replace("\\end{cases}", "\\end{array} \\end{cases}")

    replacements = {
        "\\leq": r"\\le",
        "\\geq": r"\\ge",
        "\\neq": r"\\ne",
        "\\to": r"\\rightarrow",
        "\\dfrac": r"\\frac",
        "\\tfrac": r"\\frac",
        "\\cfrac": r"\\frac",
        "\\unit{": r"\\mathrm{",
        "\\varnothing": r"\\emptyset",
        "\\geqslant": r"\\ge",
        "\\leqslant": r"\\le",
        "\\widehat": r"\\hat",
        "\\widetilde": r"\\tilde",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return re.sub(r"\s+", " ", text).strip()


def _latex_as_text(latex: str, *, display: bool) -> str:
    text = re.sub(r"\^\s*\{\s*\\prime\s*\}", "′", latex)
    text = re.sub(r"\\prime\b", "′", text)
    text = text.replace("\\infty", "∞").replace("\\pi", "π")
    text = text.replace("\\square", "ℝ").replace("\\setminus", "\\")
    text = text.replace("\\{", "{").replace("\\}", "}")
    if display:
        return f"[{text}]"
    return text


def _split_math_segments(text: str) -> list[tuple[bool, str, bool]]:
    segments: list[tuple[bool, str, bool]] = []
    index = 0
    while index < len(text):
        start = _find_next_math_start(text, index)
        if start is None:
            segments.append((False, text[index:], False))
            break
        start_index, left, right, display = start
        if start_index > index:
            segments.append((False, text[index:start_index], False))
        end_index = _find_unescaped(text, right, start_index + len(left))
        if end_index is None:
            segments.append((False, text[start_index:], False))
            break
        latex = text[start_index + len(left) : end_index].strip()
        segments.append((True, latex, display))
        index = end_index + len(right)
    return segments


def _split_implicit_latex_segments(text: str) -> list[tuple[bool, str]]:
    if not _looks_like_implicit_latex(text):
        return [(False, text)]

    segments: list[tuple[bool, str]] = []
    cursor = 0
    for match in re.finditer(r"\\[A-Za-z]+(?:\s*[_^]\s*(?:\{[^{}]*\}|[A-Za-z0-9+-]+)|\s*\{[^{}]*\}|\s*\([^()]*\)|\s*\[[^\[\]]*\])*|(?:[A-Za-z0-9]+|[xyz])(?:[_^]\s*\{[^{}]+\}|[_^]\s*[A-Za-z0-9+-]+)+", text):
        start, end = match.span()
        while start > cursor and text[start - 1] in "([{":
            start -= 1
        if text[start:end].startswith("\\int"):
            differential = re.match(r"\s*d[A-Za-z]", text[end:])
            if differential:
                end += differential.end()
        while end < len(text) and text[end] in ")]}.;,":
            if text[end] in ".;," and (end + 1 >= len(text) or not text[end + 1].isspace()):
                break
            end += 1
        if start > cursor:
            segments.append((False, text[cursor:start]))
        segments.append((True, text[start:end].strip()))
        cursor = end
    if cursor < len(text):
        segments.append((False, text[cursor:]))
    return [(is_math, value) for is_math, value in segments if value]


def _looks_like_implicit_latex(text: str) -> bool:
    return bool(re.search(r"\\(frac|sqrt|vec|overrightarrow|left|right|int|log|ln|sin|cos|tan|pi|infty|square|setminus|leq?|geq?|neq|mathbb|mathrm)\b|[A-Za-z0-9][_\^]\s*(?:\{|[A-Za-z0-9+-])", text))


def _find_next_math_start(text: str, start_at: int) -> tuple[int, str, str, bool] | None:
    candidates: list[tuple[int, str, str, bool]] = []
    for left, right, display in (("$$", "$$", True), ("\\[", "\\]", True), ("\\(", "\\)", False)):
        index = _find_unescaped(text, left, start_at)
        if index is not None:
            candidates.append((index, left, right, display))

    dollar_index = _find_inline_dollar_start(text, start_at)
    if dollar_index is not None:
        candidates.append((dollar_index, "$", "$", False))

    if not candidates:
        return None
    return min(candidates, key=lambda candidate: candidate[0])


def _find_inline_dollar_start(text: str, start_at: int) -> int | None:
    index = _find_unescaped(text, "$", start_at)
    while index is not None:
        if _is_probable_inline_math_dollar(text, index):
            return index
        index = _find_unescaped(text, "$", index + 1)
    return None


def _is_probable_inline_math_dollar(text: str, index: int) -> bool:
    if index + 1 >= len(text) or text[index + 1].isspace() or text[index + 1] == "$":
        return False
    end_index = _find_unescaped(text, "$", index + 1)
    if end_index is None or end_index == index + 1:
        return False
    if text[end_index - 1].isspace():
        return False
    content = text[index + 1 : end_index]
    return bool(re.search(r"\\[A-Za-z]+|[_^{}=+*/<>]|\d\s*[=+*/^_-]|[A-Za-z]\s*[_^=]", content))


def _find_unescaped(text: str, needle: str, start_at: int) -> int | None:
    index = text.find(needle, start_at)
    while index != -1:
        if not _is_escaped(text, index):
            return index
        index = text.find(needle, index + len(needle))
    return None


def _is_escaped(text: str, index: int) -> bool:
    slash_count = 0
    cursor = index - 1
    while cursor >= 0 and text[cursor] == "\\":
        slash_count += 1
        cursor -= 1
    return slash_count % 2 == 1


def _is_display_math_line(text: str) -> bool:
    return (
        (text.startswith("$$") and text.endswith("$$") and len(text) > 4)
        or (text.startswith("\\[") and text.endswith("\\]") and len(text) > 4)
    )


def _strip_math_delimiters(text: str) -> str:
    stripped = text.strip()
    for left, right in (("$$", "$$"), ("\\[", "\\]"), ("\\(", "\\)")):
        if stripped.startswith(left) and stripped.endswith(right) and len(stripped) > len(left) + len(right):
            return stripped[len(left) : -len(right)].strip()
    if stripped.startswith("$") and stripped.endswith("$") and _is_probable_inline_math_dollar(stripped, 0):
        return stripped[1:-1].strip()
    return stripped


def _html_table_to_matrix(html: str) -> list[list[str]]:
    if not html or "<table" not in html.lower():
        return []
    soup = BeautifulSoup(html, "html.parser")
    table = soup.find("table")
    if table is None:
        return []
    matrix: list[list[str]] = []
    for tr in table.find_all("tr"):
        row: list[str] = []
        for cell in tr.find_all(["td", "th"], recursive=False):
            text = cell.get_text(" ", strip=True)
            try:
                colspan = max(1, int(cell.get("colspan") or 1))
            except ValueError:
                colspan = 1
            row.extend([text] * colspan)
        if row:
            matrix.append(row)
    return matrix


def _markdown_table_to_html(lines: list[str]) -> str:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        if cells:
            rows.append(cells)
    if not rows:
        return ""
    html_rows = []
    for row in rows:
        cells = "".join(f"<td>{_escape_html(cell)}</td>" for cell in row)
        html_rows.append(f"<tr>{cells}</tr>")
    return f"<table>{''.join(html_rows)}</table>"


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _resolve_artifact_path(raw_path: str, base_dirs: list[Path]) -> Path | None:
    if not raw_path:
        return None
    path = Path(raw_path)
    if path.is_absolute():
        return path
    normalized = raw_path.replace("\\", "/").lstrip("/")
    filename = Path(normalized).name
    search_roots: list[Path] = []
    for base_dir in base_dirs:
        if base_dir and base_dir not in search_roots:
            search_roots.append(base_dir)
        for parent in list(base_dir.parents)[:3] if base_dir else []:
            if parent not in search_roots:
                search_roots.append(parent)
    for base_dir in search_roots:
        candidate = base_dir / normalized
        if candidate.exists():
            return candidate
        image_candidate = base_dir / "images" / filename
        if image_candidate.exists():
            return image_candidate
        matches = list(base_dir.rglob(filename))
        if matches:
            return matches[0]
    return None


def _latest_file(root: Path, pattern: str) -> Path | None:
    files = [path for path in root.rglob(pattern) if path.is_file()]
    if not files:
        return None
    return max(files, key=lambda path: path.stat().st_mtime)


def _artifact_kind(path: Path) -> str:
    name = path.name.lower()
    if "llm_review" in path.parts:
        return "llm_review"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith(".md"):
        return "markdown"
    if name.endswith("_layout.pdf"):
        return "layout"
    if name.endswith(".json"):
        return "json"
    if name.endswith(".log"):
        return "log"
    return "artifact"


def _artifact_label(path: Path) -> str:
    kind = _artifact_kind(path)
    labels = {
        "docx": "Word DOCX",
        "markdown": "Markdown",
        "layout": "Layout PDF",
        "json": "Structured JSON",
        "log": "MinerU log",
        "llm_review": "LLM Review",
    }
    return labels.get(kind, path.name)


def _is_relative_to(path: Path, parent: Path) -> bool:
    try:
        path.relative_to(parent)
        return True
    except ValueError:
        return False
